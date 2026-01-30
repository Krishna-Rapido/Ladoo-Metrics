from __future__ import annotations

import io
import os
import secrets
import tempfile
import shutil
import asyncio
from pathlib import Path
from typing import Dict, Optional, Union

import duckdb
import pandas as pd
from fastapi import Depends, FastAPI, File, Header, HTTPException, UploadFile, Request, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response, FileResponse, StreamingResponse

import schemas
from schemas import (
    ErrorResponse,
    MetricsRequest,
    MetricsResponse,
    TimeSeriesPoint,
    UploadResponse,
    FunnelRequest,
    FunnelResponse,
    FunnelPoint,
    StatTestRequest,
    StatTestResult,
    CohortAggregationResponse,
    CohortAggregationRow,
    CaptainLevelRequest,
    CaptainLevelResponse,
    CaptainLevelAggregationRow,
    MobileNumberUploadResponse,
    CaptainIdRequest,
    CaptainIdResponse,
    AOFunnelRequest,
    AOFunnelResponse,
    DaprBucketRequest,
    DaprBucketResponse,
    Fe2NetRequest,
    Fe2NetResponse,
    RtuPerformanceRequest,
    RtuPerformanceResponse,
    R2ARequest,
    R2AResponse,
    R2APercentageRequest,
    R2APercentageResponse,
    A2PhhSummaryRequest,
    A2PhhSummaryResponse,
    ReportItem,
    ReportAddRequest,
    ReportAddResponse,
    ReportUpdateCommentRequest,
    ReportUpdateTitleRequest,
    ReportListResponse,
    ReportExportResponse,
    InsightsRequest,
    InsightsResponse,
    InsightsTimeSeriesPoint,
    InsightsSummaryRow,
    FunctionTestRequest,
    FunctionTestResponse,
    FunctionExecuteRequest,
    FunctionExecuteResponse,
    FunctionPreviewRequest,
    FunctionPreviewResponse,
    FunctionJoinRequest,
    FunctionJoinResponse,
    FunctionTemplateResponse,
    PivotRequest,
    PivotResponse,
)
from function_executor import (
    test_function,
    execute_function,
    join_with_csv,
    FUNCTION_TEMPLATE,
)
from transformations import (
    aggregate_time_series,
    filter_by_date_range,
    normalized_growth,
    rolling_average,
    subset_by_cohorts,
    get_cohort,
    compute_cohort_funnel_timeseries,
    compute_metric_timeseries_by_cohort,
)


app = FastAPI(title="Cohort Metrics API")

# Configuration for large file uploads (5GB max)
MAX_FILE_SIZE = 5 * 1024 * 1024 * 1024  # 5GB in bytes
CHUNK_SIZE = 1024 * 1024  # 1MB chunks for reading
CSV_CHUNK_SIZE = 100_000  # 100k rows per chunk for pandas

# Threshold for using DuckDB-based storage (files larger than this use parquet + DuckDB)
DUCKDB_THRESHOLD = 50 * 1024 * 1024  # 50MB

# Session store - maps session_id to either:
#   - pandas DataFrame (for small files, backward compatibility)
#   - dict with parquet_path and metadata (for large files, DuckDB-based)
SESSION_STORE: Dict[str, Union[pd.DataFrame, dict]] = {}

# Session file store - maps session_id to temp file path (for large files)
SESSION_FILE_STORE: Dict[str, str] = {}

# Upload progress store - maps upload_id to progress info
UPLOAD_PROGRESS: Dict[str, dict] = {}

# Funnel analysis session store - separate from main session store
FUNNEL_SESSION_STORE: Dict[str, pd.DataFrame] = {}

# Report builder session store - maps report_id to list of report items
REPORT_STORE: Dict[str, list[dict]] = {}

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://internal-tools-v1-1.onrender.com",  # Your frontend URL
                    "http://localhost:5173", 
                    "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def is_duckdb_session(session_id: str) -> bool:
    """Check if a session uses DuckDB-based storage (parquet file on disk)."""
    if session_id not in SESSION_STORE:
        return False
    session = SESSION_STORE[session_id]
    return isinstance(session, dict) and "parquet_path" in session


def get_session_parquet_path(x_session_id: Optional[str] = Header(default=None)) -> str:
    """Get the parquet file path for a DuckDB-based session."""
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    
    session = SESSION_STORE[x_session_id]
    if isinstance(session, dict) and "parquet_path" in session:
        return session["parquet_path"]
    
    raise HTTPException(status_code=400, detail="Session does not use DuckDB storage")


def get_session_metadata(x_session_id: Optional[str] = Header(default=None)) -> dict:
    """Get session metadata (works for both pandas and DuckDB sessions)."""
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    
    session = SESSION_STORE[x_session_id]
    if isinstance(session, dict) and "parquet_path" in session:
        return session
    
    # For pandas DataFrame sessions, extract metadata
    df = session
    return {
        "num_rows": len(df),
        "columns": list(df.columns),
        "cohorts": sorted(df["cohort"].dropna().unique().tolist()) if "cohort" in df.columns else [],
        "date_min": df["date"].min().strftime("%Y-%m-%d") if "date" in df.columns else None,
        "date_max": df["date"].max().strftime("%Y-%m-%d") if "date" in df.columns else None,
        "metrics": [c for c in df.columns if c not in {"cohort", "date", "time"}],
    }


def query_session_duckdb(session_id: str, sql: str) -> pd.DataFrame:
    """
    Execute SQL query on session's Parquet file using DuckDB.
    The query should reference the table as 'data'.
    Returns results as a pandas DataFrame.
    """
    if session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid session_id")
    
    session = SESSION_STORE[session_id]
    if not isinstance(session, dict) or "parquet_path" not in session:
        raise HTTPException(status_code=400, detail="Session does not use DuckDB storage")
    
    parquet_path = session["parquet_path"]
    
    if not os.path.exists(parquet_path):
        raise HTTPException(status_code=400, detail="Session data file not found")
    
    # Execute query using DuckDB
    con = duckdb.connect()
    try:
        # Register the parquet file as a table named 'data'
        con.execute(f"CREATE VIEW data AS SELECT * FROM read_parquet('{parquet_path}')")
        result = con.execute(sql).fetchdf()
        return result
    finally:
        con.close()


def get_session_df(x_session_id: Optional[str] = Header(default=None)) -> pd.DataFrame:
    """
    Get session DataFrame. For DuckDB sessions, loads the full data into pandas.
    WARNING: For large files, prefer using query_session_duckdb() with specific queries.
    """
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    
    session = SESSION_STORE[x_session_id]
    
    # If it's already a DataFrame, return it
    if isinstance(session, pd.DataFrame):
        return session
    
    # If it's a DuckDB session, load from parquet
    if isinstance(session, dict) and "parquet_path" in session:
        parquet_path = session["parquet_path"]
        if not os.path.exists(parquet_path):
            raise HTTPException(status_code=400, detail="Session data file not found")
        
        # Use DuckDB to read the parquet file (more memory efficient than pyarrow for large files)
        con = duckdb.connect()
        try:
            df = con.execute(f"SELECT * FROM read_parquet('{parquet_path}')").fetchdf()
            return df
        finally:
            con.close()
    
    raise HTTPException(status_code=400, detail="Invalid session data format")


async def stream_file_to_disk(file: UploadFile, temp_path: str, upload_id: str = None) -> int:
    """Stream uploaded file to disk in chunks to avoid memory issues."""
    total_size = 0
    with open(temp_path, 'wb') as f:
        while True:
            chunk = await file.read(CHUNK_SIZE)
            if not chunk:
                break
            f.write(chunk)
            total_size += len(chunk)
            
            # Update progress if upload_id provided
            if upload_id and upload_id in UPLOAD_PROGRESS:
                UPLOAD_PROGRESS[upload_id]['bytes_uploaded'] = total_size
            
            # Check file size limit
            if total_size > MAX_FILE_SIZE:
                os.unlink(temp_path)
                raise HTTPException(status_code=413, detail=f"File too large. Maximum size is {MAX_FILE_SIZE / (1024**3):.1f}GB")
    
    return total_size


def detect_csv_delimiter(file_path: str) -> str:
    """Detect CSV delimiter by reading first few lines."""
    import csv
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        # Read first 8KB to detect delimiter
        sample = f.read(8192)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=',;\t|')
            return dialect.delimiter
        except csv.Error:
            return ','  # Default to comma


def process_csv_to_parquet(temp_path: str, parquet_path: str) -> dict:
    """
    Process CSV file in chunks and write to Parquet for DuckDB-based storage.
    Returns metadata dict with parquet_path, num_rows, columns, cohorts, date range, metrics.
    Does NOT load the full data into memory.
    """
    import pyarrow as pa
    import pyarrow.parquet as pq
    
    # Detect delimiter first
    delimiter = detect_csv_delimiter(temp_path)
    
    # First, read a small sample to validate structure
    try:
        sample_df = pd.read_csv(temp_path, nrows=100, sep=delimiter, encoding='utf-8', encoding_errors='replace')
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to read CSV: {exc}")
    
    # Validate required columns
    if "cohort" not in sample_df.columns:
        raise HTTPException(status_code=400, detail="CSV missing required column: 'cohort'")
    if ("date" not in sample_df.columns) and ("time" not in sample_df.columns):
        raise HTTPException(status_code=400, detail="CSV must include 'date' (YYYY-MM-DD) or 'time' (YYYYMMDD)")
    
    # Determine dtypes for efficient reading
    has_date = "date" in sample_df.columns
    columns_list = list(sample_df.columns)
    
    # Determine columns that should be numeric
    numeric_cols = []
    for c in sample_df.columns:
        if c not in {"cohort", "date", "time"}:
            try:
                pd.to_numeric(sample_df[c], errors='raise')
                numeric_cols.append(c)
            except (ValueError, TypeError):
                pass
    
    # Process CSV in chunks and write to Parquet incrementally
    total_rows = 0
    invalid_date_count = 0
    writer = None
    
    try:
        for chunk_idx, chunk in enumerate(pd.read_csv(
            temp_path, 
            chunksize=CSV_CHUNK_SIZE, 
            low_memory=False,
            sep=delimiter,
            encoding='utf-8',
            encoding_errors='replace',
            on_bad_lines='warn'
        )):
            # Process date column
            if has_date:
                chunk["date"] = pd.to_datetime(chunk["date"], errors="coerce")
            else:
                chunk["date"] = pd.to_datetime(chunk["time"].astype(str), format="%Y%m%d", errors="coerce")
            
            # Count invalid dates
            invalid_date_count += chunk["date"].isna().sum()
            
            # Process cohort column
            chunk["cohort"] = chunk["cohort"].astype(str)
            
            # Numeric coercion for metric columns
            for c in numeric_cols:
                if c in chunk.columns:
                    chunk[c] = pd.to_numeric(chunk[c], errors="coerce")
            
            # Convert to pyarrow table
            table = pa.Table.from_pandas(chunk, preserve_index=False)
            
            # Write to parquet (append mode)
            if writer is None:
                writer = pq.ParquetWriter(parquet_path, table.schema)
            writer.write_table(table)
            
            total_rows += len(chunk)
            
            # Free memory
            del chunk
            del table
        
        if writer:
            writer.close()
            writer = None
        
        # Check for invalid dates (warn but don't fail for large files)
        if invalid_date_count > 0:
            # For large files, log warning but allow if less than 1% invalid
            invalid_pct = (invalid_date_count / total_rows) * 100
            if invalid_pct > 1:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Too many invalid date/time values: {invalid_date_count} ({invalid_pct:.1f}%)"
                )
        
        # Use DuckDB to extract metadata WITHOUT loading full data into memory
        con = duckdb.connect()
        try:
            # Get date range and row count
            meta = con.execute(f"""
                SELECT 
                    COUNT(*) as num_rows,
                    MIN(date) as date_min,
                    MAX(date) as date_max
                FROM read_parquet('{parquet_path}')
            """).fetchone()
            
            # Get unique cohorts
            cohorts_result = con.execute(f"""
                SELECT DISTINCT cohort 
                FROM read_parquet('{parquet_path}') 
                WHERE cohort IS NOT NULL
                ORDER BY cohort
            """).fetchall()
            
            cohorts = [c[0] for c in cohorts_result]
            
            # Format dates
            date_min = meta[1]
            date_max = meta[2]
            if hasattr(date_min, 'strftime'):
                date_min_str = date_min.strftime("%Y-%m-%d")
            else:
                date_min_str = str(date_min)[:10] if date_min else None
                
            if hasattr(date_max, 'strftime'):
                date_max_str = date_max.strftime("%Y-%m-%d")
            else:
                date_max_str = str(date_max)[:10] if date_max else None
            
            # Get column info from parquet schema
            schema_result = con.execute(f"DESCRIBE SELECT * FROM read_parquet('{parquet_path}')").fetchall()
            columns = [row[0] for row in schema_result]
            
            # Determine metric columns (exclude cohort, date, time)
            metrics = [c for c in columns if c not in {"cohort", "date", "time"}]
            
            # Determine numeric columns
            numeric_columns = []
            categorical_columns = []
            for row in schema_result:
                col_name, col_type = row[0], row[1]
                if col_name not in {"cohort", "date", "time"}:
                    if any(t in col_type.lower() for t in ['int', 'float', 'double', 'decimal', 'numeric']):
                        numeric_columns.append(col_name)
                    elif col_type.lower() in ['varchar', 'string']:
                        categorical_columns.append(col_name)
            
            return {
                "parquet_path": parquet_path,
                "num_rows": meta[0],
                "columns": columns,
                "cohorts": cohorts,
                "date_min": date_min_str,
                "date_max": date_max_str,
                "metrics": metrics,
                "numeric_columns": numeric_columns,
                "categorical_columns": categorical_columns,
                "invalid_dates": invalid_date_count,
            }
        finally:
            con.close()
        
    except HTTPException:
        raise
    except Exception as exc:
        # Clean up on error
        if writer:
            try:
                writer.close()
            except Exception:
                pass
        if os.path.exists(parquet_path):
            os.unlink(parquet_path)
        raise HTTPException(status_code=400, detail=f"Failed to process CSV: {exc}")


def process_csv_chunked(temp_path: str, session_id: str) -> tuple[pd.DataFrame, dict]:
    """
    Legacy function for backward compatibility.
    Process CSV file in chunks and return DataFrame.
    For large files, prefer using process_csv_to_parquet() which doesn't load into memory.
    """
    import pyarrow.parquet as pq
    
    # Create temp parquet path
    parquet_path = temp_path.replace('.csv', '.parquet')
    
    # Process to parquet
    metadata = process_csv_to_parquet(temp_path, parquet_path)
    
    # Read back as DataFrame (for backward compatibility with small file handling)
    df = pq.read_table(parquet_path).to_pandas()
    
    # Clean up parquet file since we loaded into memory
    if os.path.exists(parquet_path):
        os.unlink(parquet_path)
    
    return df, {"total_rows": metadata["num_rows"], "invalid_dates": metadata["invalid_dates"]}


@app.post("/upload", response_model=UploadResponse, responses={400: {"model": ErrorResponse}})
async def upload_csv(file: UploadFile = File(...)) -> UploadResponse:
    """
    Upload a CSV file for analysis. Supports files up to 5GB.
    
    For very large files (>100MB), consider using /upload/chunked endpoint
    which provides progress tracking.
    """
    if not file.filename.lower().endswith(".csv"):
        raise HTTPException(status_code=400, detail="Only CSV files are supported")
    
    session_id = secrets.token_hex(16)
    
    # Create temp file for streaming
    temp_dir = tempfile.mkdtemp()
    temp_path = os.path.join(temp_dir, "upload.csv")
    
    try:
        # Stream file to disk instead of loading into memory
        file_size = await stream_file_to_disk(file, temp_path)
        
        # Detect delimiter
        delimiter = detect_csv_delimiter(temp_path)
        
        # For smaller files (<DUCKDB_THRESHOLD), use the traditional in-memory approach for speed
        if file_size < DUCKDB_THRESHOLD:
            df = pd.read_csv(temp_path, sep=delimiter, encoding='utf-8', encoding_errors='replace')
            
            # Validate and process
            if "cohort" not in df.columns:
                raise HTTPException(status_code=400, detail="CSV missing required column: 'cohort'")
            if ("date" not in df.columns) and ("time" not in df.columns):
                raise HTTPException(status_code=400, detail="CSV must include 'date' (YYYY-MM-DD) or 'time' (YYYYMMDD)")
            
            df = df.copy()
            if "date" in df.columns:
                df["date"] = pd.to_datetime(df["date"], errors="coerce")
            else:
                df["date"] = pd.to_datetime(df["time"].astype(str), format="%Y%m%d", errors="coerce")
            
            if df["date"].isna().any():
                x = df["date"].isna().sum()
                raise HTTPException(status_code=400, detail=f"Invalid date/time values found = {x}")
            
            df["cohort"] = df["cohort"].astype(str)
            
            # Best-effort numeric coercion
            for c in df.columns:
                if c in {"cohort", "date", "time"}:
                    continue
                try:
                    df[c] = pd.to_numeric(df[c], errors="ignore")
                except Exception:
                    pass
            
            # Store DataFrame in session (in-memory for small files)
            SESSION_STORE[session_id] = df
            SESSION_FILE_STORE[session_id] = temp_path
            
            cohorts = sorted(df["cohort"].dropna().unique().tolist())
            date_min = df["date"].min()
            date_max = df["date"].max()
            metric_candidates = [c for c in df.columns if c not in {"cohort", "date", "time"}]
            
            return UploadResponse(
                session_id=session_id,
                num_rows=df.shape[0],
                columns=list(df.columns.astype(str)),
                cohorts=cohorts,
                date_min=date_min.strftime("%Y-%m-%d"),
                date_max=date_max.strftime("%Y-%m-%d"),
                metrics=metric_candidates,
            )
        else:
            # For large files, use DuckDB-based storage (parquet on disk)
            parquet_path = os.path.join(temp_dir, "data.parquet")
            
            # Process CSV to Parquet and get metadata (without loading into memory)
            session_metadata = process_csv_to_parquet(temp_path, parquet_path)
            
            # Store metadata in session (NOT the full DataFrame)
            SESSION_STORE[session_id] = session_metadata
            SESSION_FILE_STORE[session_id] = temp_dir  # Store dir for cleanup
            
            return UploadResponse(
                session_id=session_id,
                num_rows=session_metadata["num_rows"],
                columns=session_metadata["columns"],
                cohorts=session_metadata["cohorts"],
                date_min=session_metadata["date_min"],
                date_max=session_metadata["date_max"],
                metrics=session_metadata["metrics"],
            )
    except HTTPException:
        # Clean up temp dir on validation errors
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise
    except Exception as exc:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise HTTPException(status_code=400, detail=f"Failed to process CSV: {exc}")


@app.post("/upload/init")
async def init_chunked_upload(request: Request):
    """
    Initialize a chunked upload session. Returns an upload_id for tracking progress.
    Use this for large files (>100MB) to enable progress tracking.
    """
    body = await request.json()
    filename = body.get("filename", "upload.csv")
    file_size = body.get("file_size", 0)
    
    if not filename.lower().endswith(".csv"):
        raise HTTPException(status_code=400, detail="Only CSV files are supported")
    
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File too large. Maximum size is {MAX_FILE_SIZE / (1024**3):.1f}GB")
    
    upload_id = secrets.token_hex(16)
    temp_dir = tempfile.mkdtemp()
    temp_path = os.path.join(temp_dir, "upload.csv")
    
    UPLOAD_PROGRESS[upload_id] = {
        "status": "initialized",
        "filename": filename,
        "total_size": file_size,
        "bytes_uploaded": 0,
        "temp_path": temp_path,
        "temp_dir": temp_dir,
    }
    
    return {"upload_id": upload_id, "max_chunk_size": CHUNK_SIZE}


@app.post("/upload/chunk/{upload_id}")
async def upload_chunk(upload_id: str, file: UploadFile = File(...)):
    """
    Upload a chunk of the file. Chunks should be uploaded sequentially.
    """
    if upload_id not in UPLOAD_PROGRESS:
        raise HTTPException(status_code=404, detail="Upload session not found")
    
    progress = UPLOAD_PROGRESS[upload_id]
    temp_path = progress["temp_path"]
    
    # Read chunk and append to file
    chunk_data = await file.read()
    
    with open(temp_path, 'ab') as f:
        f.write(chunk_data)
    
    progress["bytes_uploaded"] += len(chunk_data)
    progress["status"] = "uploading"
    
    return {
        "upload_id": upload_id,
        "bytes_uploaded": progress["bytes_uploaded"],
        "total_size": progress["total_size"],
        "progress": progress["bytes_uploaded"] / max(progress["total_size"], 1) * 100
    }


@app.get("/upload/progress/{upload_id}")
async def get_upload_progress(upload_id: str):
    """Get the progress of an ongoing upload."""
    if upload_id not in UPLOAD_PROGRESS:
        raise HTTPException(status_code=404, detail="Upload session not found")
    
    progress = UPLOAD_PROGRESS[upload_id]
    return {
        "upload_id": upload_id,
        "status": progress["status"],
        "bytes_uploaded": progress["bytes_uploaded"],
        "total_size": progress["total_size"],
        "progress": progress["bytes_uploaded"] / max(progress["total_size"], 1) * 100
    }


@app.post("/upload/complete/{upload_id}", response_model=UploadResponse)
async def complete_chunked_upload(upload_id: str):
    """
    Complete a chunked upload and process the CSV file.
    """
    if upload_id not in UPLOAD_PROGRESS:
        raise HTTPException(status_code=404, detail="Upload session not found")
    
    progress = UPLOAD_PROGRESS[upload_id]
    temp_path = progress["temp_path"]
    temp_dir = progress["temp_dir"]
    
    progress["status"] = "processing"
    session_id = secrets.token_hex(16)
    
    try:
        file_size = progress["bytes_uploaded"]
        
        # Detect delimiter
        delimiter = detect_csv_delimiter(temp_path)
        
        # For smaller files (<DUCKDB_THRESHOLD), use in-memory approach
        if file_size < DUCKDB_THRESHOLD:
            df = pd.read_csv(temp_path, sep=delimiter, encoding='utf-8', encoding_errors='replace')
            
            if "cohort" not in df.columns:
                raise HTTPException(status_code=400, detail="CSV missing required column: 'cohort'")
            if ("date" not in df.columns) and ("time" not in df.columns):
                raise HTTPException(status_code=400, detail="CSV must include 'date' (YYYY-MM-DD) or 'time' (YYYYMMDD)")
            
            df = df.copy()
            if "date" in df.columns:
                df["date"] = pd.to_datetime(df["date"], errors="coerce")
            else:
                df["date"] = pd.to_datetime(df["time"].astype(str), format="%Y%m%d", errors="coerce")
            
            if df["date"].isna().any():
                x = df["date"].isna().sum()
                raise HTTPException(status_code=400, detail=f"Invalid date/time values found = {x}")
            
            df["cohort"] = df["cohort"].astype(str)
            
            for c in df.columns:
                if c in {"cohort", "date", "time"}:
                    continue
                try:
                    df[c] = pd.to_numeric(df[c], errors="ignore")
                except Exception:
                    pass
            
            SESSION_STORE[session_id] = df
            SESSION_FILE_STORE[session_id] = temp_path
            
            cohorts = sorted(df["cohort"].dropna().unique().tolist())
            date_min = df["date"].min()
            date_max = df["date"].max()
            metric_candidates = [c for c in df.columns if c not in {"cohort", "date", "time"}]
            
            progress["status"] = "completed"
            progress["session_id"] = session_id
            
            return UploadResponse(
                session_id=session_id,
                num_rows=df.shape[0],
                columns=list(df.columns.astype(str)),
                cohorts=cohorts,
                date_min=date_min.strftime("%Y-%m-%d"),
                date_max=date_max.strftime("%Y-%m-%d"),
                metrics=metric_candidates,
            )
        else:
            # For large files, use DuckDB-based storage (parquet on disk)
            parquet_path = os.path.join(temp_dir, "data.parquet")
            
            # Process CSV to Parquet and get metadata (without loading into memory)
            session_metadata = process_csv_to_parquet(temp_path, parquet_path)
            
            # Store metadata in session (NOT the full DataFrame)
            SESSION_STORE[session_id] = session_metadata
            SESSION_FILE_STORE[session_id] = temp_dir  # Store dir for cleanup
            
            progress["status"] = "completed"
            progress["session_id"] = session_id
            
            return UploadResponse(
                session_id=session_id,
                num_rows=session_metadata["num_rows"],
                columns=session_metadata["columns"],
                cohorts=session_metadata["cohorts"],
                date_min=session_metadata["date_min"],
                date_max=session_metadata["date_max"],
                metrics=session_metadata["metrics"],
            )
    except HTTPException:
        progress["status"] = "error"
        shutil.rmtree(temp_dir, ignore_errors=True)
        del UPLOAD_PROGRESS[upload_id]
        raise
    except Exception as exc:
        progress["status"] = "error"
        shutil.rmtree(temp_dir, ignore_errors=True)
        del UPLOAD_PROGRESS[upload_id]
        raise HTTPException(status_code=400, detail=f"Failed to process CSV: {exc}")


@app.delete("/upload/{upload_id}")
async def cancel_upload(upload_id: str):
    """Cancel an ongoing upload and clean up resources."""
    if upload_id not in UPLOAD_PROGRESS:
        raise HTTPException(status_code=404, detail="Upload session not found")
    
    progress = UPLOAD_PROGRESS[upload_id]
    temp_dir = progress.get("temp_dir")
    
    if temp_dir:
        shutil.rmtree(temp_dir, ignore_errors=True)
    
    del UPLOAD_PROGRESS[upload_id]
    
    return {"status": "cancelled", "upload_id": upload_id}


@app.get("/meta")
def get_meta(x_session_id: Optional[str] = Header(default=None)):
    """Get metadata for the session - works with both DuckDB and pandas sessions."""
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    
    session = SESSION_STORE[x_session_id]
    
    # If it's a DuckDB session (dict with parquet_path), return stored metadata
    if isinstance(session, dict) and "parquet_path" in session:
        # Get categorical columns from DuckDB if not already cached
        categorical_columns = session.get("categorical_columns", [])
        
        return {
            "cohorts": session["cohorts"],
            "date_min": session["date_min"],
            "date_max": session["date_max"],
            "metrics": session["metrics"],
            "categorical_columns": sorted(categorical_columns)
        }
    
    # For pandas DataFrame sessions, compute metadata
    df = session
    cohorts = sorted(df["cohort"].dropna().unique().tolist())
    date_min = df["date"].min().strftime("%Y-%m-%d")
    date_max = df["date"].max().strftime("%Y-%m-%d")
    metrics = [c for c in df.columns if c not in {"cohort", "date", "time"}]
    
    # Identify categorical columns (non-numeric columns excluding cohort and date)
    categorical_columns = []
    for col in df.columns:
        if col not in {"cohort", "date", "time"}:
            # Check if column is categorical (object type or has limited unique values)
            if df[col].dtype == 'object' or df[col].dtype.name == 'category':
                categorical_columns.append(col)
            elif df[col].dtype in ['int64', 'int32', 'float64', 'float32']:
                # If numeric but has limited unique values (less than 20), consider it categorical
                unique_count = df[col].nunique()
                if unique_count < 20 and unique_count < len(df) * 0.1:
                    categorical_columns.append(col)
    
    return {
        "cohorts": cohorts,
        "date_min": date_min,
        "date_max": date_max,
        "metrics": metrics,
        "categorical_columns": sorted(categorical_columns)
    }


@app.get("/insights")
def insights_help():
    """Helpful response for browser hits; the actual Insights compute route is POST /insights."""
    return {
        "ok": True,
        "message": "Use POST /insights with JSON body + x-session-id header (set by /upload).",
        "example": {
            "method": "POST",
            "path": "/insights",
            "headers": {"x-session-id": "<session_id>"},
            "json": {
                "test_cohort": "<test cohort>",
                "control_cohort": "<control cohort>",
                "pre_period": {"start_date": "2024-01-01", "end_date": "2024-02-14"},
                "post_period": {"start_date": "2024-02-15", "end_date": "2024-03-31"},
                "metrics": [{"column": "<metric column>", "agg_func": "sum"}],
            },
        },
    }


def compute_insights_duckdb(payload: InsightsRequest, parquet_path: str) -> InsightsResponse:
    """
    Compute insights using DuckDB queries on parquet file.
    Memory efficient - doesn't load full data into memory.
    """
    if not payload.metrics:
        raise HTTPException(status_code=400, detail="No metrics selected")
    
    if not payload.test_cohort or not payload.control_cohort:
        raise HTTPException(status_code=400, detail="Both test_cohort and control_cohort are required")
    
    con = duckdb.connect()
    try:
        # Get column types from parquet so we use correct aggregation for numeric vs string columns
        type_df = con.execute(f"DESCRIBE SELECT * FROM read_parquet('{parquet_path}')").fetchdf()
        # DuckDB DESCRIBE may return 'column_name'/'column_type' or 'Column'/'Type'
        name_col = type_df.columns[0]
        type_col = type_df.columns[1]
        col_types = dict(zip(type_df[name_col], type_df[type_col]))
        
        def is_numeric_col(col: str) -> bool:
            t = col_types.get(col, "")
            t_lower = str(t).lower()
            return any(x in t_lower for x in ["int", "float", "double", "decimal", "numeric", "bigint", "smallint"])
        
        # Quote column names for DuckDB (handles special characters like ':')
        def quote_col(col: str) -> str:
            # Escape any double quotes in the column name and wrap in double quotes
            return f'"{col.replace(chr(34), chr(34)+chr(34))}"'
        
        # captain_id column (quoted) for sum_per_captain; fallback if missing
        captain_id_quoted = quote_col("captain_id") if "captain_id" in col_types else None

        # Map aggregation functions to SQL; for non-numeric columns use COUNT/COUNT DISTINCT only
        def get_sql_agg(agg: str, col: str) -> str:
            quoted = quote_col(col)
            numeric = is_numeric_col(col)
            
            if agg == "count":
                return f"COUNT({quoted})"
            if agg == "nunique":
                return f"COUNT(DISTINCT {quoted})"
            # sum_per_captain = sum(metric) / count(distinct captain_id)
            if agg == "sum_per_captain":
                if captain_id_quoted is None:
                    # No captain_id column: fall back to sum
                    if numeric:
                        return f"SUM(CAST({quoted} AS DOUBLE))"
                    return f"COUNT({quoted})"
                if numeric:
                    return f"SUM(CAST({quoted} AS DOUBLE)) / NULLIF(COUNT(DISTINCT {captain_id_quoted}), 0)"
                return f"COUNT({quoted}) / NULLIF(COUNT(DISTINCT {captain_id_quoted}), 0)"
            # For sum/mean/median on non-numeric columns, use COUNT (interpret as row count)
            if not numeric:
                return f"COUNT({quoted})"
            if agg == "sum":
                return f"SUM(CAST({quoted} AS DOUBLE))"
            if agg == "mean":
                return f"AVG(CAST({quoted} AS DOUBLE))"
            if agg == "median":
                return f"MEDIAN(CAST({quoted} AS DOUBLE))"
            raise ValueError(f"Unsupported aggregation: {agg}")
        
        # Resolve metric SQL: column may be a physical column or a ratio key "num2denom" (value = SUM(denom)/SUM(num))
        def get_metric_sql_and_agg(spec):
            col = spec.column
            agg = spec.agg_func
            if col in col_types:
                return get_sql_agg(agg, col), col, agg
            if "2" in col:
                parts = col.split("2", 1)
                if len(parts) == 2:
                    num, denom = parts[0].strip(), parts[1].strip()
                    if num in col_types and denom in col_types:
                        qn, qd = quote_col(num), quote_col(denom)
                        ratio_sql = f"SUM(CAST({qd} AS DOUBLE)) / NULLIF(SUM(CAST({qn} AS DOUBLE)), 0)"
                        return ratio_sql, col, "ratio"
            return None, None, None
        
        # Determine date range for timeseries
        start_candidates = []
        end_candidates = []
        for period in [payload.pre_period, payload.post_period]:
            if period and period.start_date:
                start_candidates.append(period.start_date)
            if period and period.end_date:
                end_candidates.append(period.end_date)
        ts_start = min(start_candidates) if start_candidates else None
        ts_end = max(end_candidates) if end_candidates else None
        
        # Build date filter clause
        date_filter = ""
        if ts_start and ts_end:
            date_filter = f"AND date >= '{ts_start}' AND date <= '{ts_end}'"
        elif ts_start:
            date_filter = f"AND date >= '{ts_start}'"
        elif ts_end:
            date_filter = f"AND date <= '{ts_end}'"
        
        # Period (pre/post) for 4-line chart
        pre_start = payload.pre_period.start_date if payload.pre_period else None
        pre_end = payload.pre_period.end_date if payload.pre_period else None
        post_start = payload.post_period.start_date if payload.post_period else None
        post_end = payload.post_period.end_date if payload.post_period else None
        period_case = "NULL"
        if pre_start and pre_end and post_start and post_end:
            period_case = f"""CASE
                WHEN date >= '{pre_start}' AND date <= '{pre_end}' THEN 'pre'
                WHEN date >= '{post_start}' AND date <= '{post_end}' THEN 'post'
                ELSE NULL
            END"""
        
        # Optional series breakout (categorical column)
        breakout_col = payload.series_breakout if payload.series_breakout and payload.series_breakout in col_types else None
        breakout_select = ""
        breakout_group = ""
        if breakout_col:
            bq = quote_col(breakout_col)
            breakout_select = f", COALESCE(CAST({bq} AS VARCHAR), 'Unknown') as breakout_value"
            breakout_group = f", breakout_value"
        
        # -------- Timeseries query --------
        time_series = []
        for spec in payload.metrics:
            sql_agg, col, agg = get_metric_sql_and_agg(spec)
            if sql_agg is None:
                continue
            ts_query = f"""
                SELECT 
                    CAST(date AS DATE) as date,
                    CASE 
                        WHEN cohort = '{payload.test_cohort}' THEN 'test' 
                        ELSE 'control' 
                    END as cohort_type,
                    {period_case} as period,
                    '{col}' as metric,
                    '{agg}' as agg_func,
                    {sql_agg} as value
                    {breakout_select}
                FROM read_parquet('{parquet_path}')
                WHERE cohort IN ('{payload.test_cohort}', '{payload.control_cohort}')
                {date_filter}
                GROUP BY CAST(date AS DATE), cohort_type, period
                    {breakout_group}
                ORDER BY date, cohort_type, period
                    {breakout_group.replace(', breakout_value', '')}
            """
            # If no breakout, remove the extra comma from GROUP BY
            if not breakout_col:
                ts_query = f"""
                SELECT 
                    CAST(date AS DATE) as date,
                    CASE 
                        WHEN cohort = '{payload.test_cohort}' THEN 'test' 
                        ELSE 'control' 
                    END as cohort_type,
                    {period_case} as period,
                    '{col}' as metric,
                    '{agg}' as agg_func,
                    {sql_agg} as value
                FROM read_parquet('{parquet_path}')
                WHERE cohort IN ('{payload.test_cohort}', '{payload.control_cohort}')
                {date_filter}
                GROUP BY CAST(date AS DATE), cohort_type, period
                ORDER BY date, cohort_type, period
            """
            else:
                ts_query = f"""
                SELECT 
                    CAST(date AS DATE) as date,
                    CASE 
                        WHEN cohort = '{payload.test_cohort}' THEN 'test' 
                        ELSE 'control' 
                    END as cohort_type,
                    {period_case} as period,
                    '{col}' as metric,
                    '{agg}' as agg_func,
                    {sql_agg} as value,
                    COALESCE(CAST({quote_col(breakout_col)} AS VARCHAR), 'Unknown') as breakout_value
                FROM read_parquet('{parquet_path}')
                WHERE cohort IN ('{payload.test_cohort}', '{payload.control_cohort}')
                {date_filter}
                GROUP BY CAST(date AS DATE), cohort_type, period, breakout_value
                ORDER BY date, cohort_type, period, breakout_value
            """
            ts_df = con.execute(ts_query).fetchdf()
            for _, row in ts_df.iterrows():
                time_series.append(InsightsTimeSeriesPoint(
                    date=str(row["date"])[:10],
                    cohort_type=str(row["cohort_type"]),
                    period=str(row["period"]) if row.get("period") is not None and str(row["period"]) != "None" else None,
                    metric=col,
                    agg_func=agg,
                    value=float(row["value"]) if pd.notna(row["value"]) else 0.0,
                    breakout_value=str(row["breakout_value"]) if breakout_col and row.get("breakout_value") is not None else None,
                ))
        
        # -------- Executive summary query --------
        summary = []
        for spec in payload.metrics:
            sql_agg, col, agg = get_metric_sql_and_agg(spec)
            if sql_agg is None:
                continue
            # Build period filters
            pre_filter = ""
            post_filter = ""
            if payload.pre_period:
                if payload.pre_period.start_date:
                    pre_filter += f" AND date >= '{payload.pre_period.start_date}'"
                if payload.pre_period.end_date:
                    pre_filter += f" AND date <= '{payload.pre_period.end_date}'"
            if payload.post_period:
                if payload.post_period.start_date:
                    post_filter += f" AND date >= '{payload.post_period.start_date}'"
                if payload.post_period.end_date:
                    post_filter += f" AND date <= '{payload.post_period.end_date}'"
            
            # Query for pre and post values per cohort_type
            summary_query = f"""
                WITH pre_data AS (
                    SELECT 
                        CASE WHEN cohort = '{payload.test_cohort}' THEN 'test' ELSE 'control' END as cohort_type,
                        {sql_agg} as value
                    FROM read_parquet('{parquet_path}')
                    WHERE cohort IN ('{payload.test_cohort}', '{payload.control_cohort}')
                    {pre_filter}
                    GROUP BY cohort_type
                ),
                post_data AS (
                    SELECT 
                        CASE WHEN cohort = '{payload.test_cohort}' THEN 'test' ELSE 'control' END as cohort_type,
                        {sql_agg} as value
                    FROM read_parquet('{parquet_path}')
                    WHERE cohort IN ('{payload.test_cohort}', '{payload.control_cohort}')
                    {post_filter}
                    GROUP BY cohort_type
                )
                SELECT 
                    p.cohort_type,
                    COALESCE(pre.value, 0) as pre_value,
                    COALESCE(post.value, 0) as post_value
                FROM (SELECT 'test' as cohort_type UNION SELECT 'control') p
                LEFT JOIN pre_data pre ON p.cohort_type = pre.cohort_type
                LEFT JOIN post_data post ON p.cohort_type = post.cohort_type
            """
            
            summary_df = con.execute(summary_query).fetchdf()
            
            control_pre = 0.0
            control_post = 0.0
            test_pre = 0.0
            test_post = 0.0
            
            for _, row in summary_df.iterrows():
                if row["cohort_type"] == "control":
                    control_pre = float(row["pre_value"]) if pd.notna(row["pre_value"]) else 0.0
                    control_post = float(row["post_value"]) if pd.notna(row["post_value"]) else 0.0
                elif row["cohort_type"] == "test":
                    test_pre = float(row["pre_value"]) if pd.notna(row["pre_value"]) else 0.0
                    test_post = float(row["post_value"]) if pd.notna(row["post_value"]) else 0.0
            
            control_delta = control_post - control_pre
            test_delta = test_post - test_pre
            diff_in_diff = test_delta - control_delta
            
            control_delta_pct = (control_delta / control_pre * 100.0) if control_pre != 0 else None
            test_delta_pct = (test_delta / test_pre * 100.0) if test_pre != 0 else None
            diff_in_diff_pct = (diff_in_diff / control_pre * 100.0) if control_pre != 0 else None
            
            summary.append(InsightsSummaryRow(
                metric=col,
                agg_func=agg,
                control_pre=control_pre,
                control_post=control_post,
                control_delta=control_delta,
                control_delta_pct=control_delta_pct,
                test_pre=test_pre,
                test_post=test_post,
                test_delta=test_delta,
                test_delta_pct=test_delta_pct,
                diff_in_diff=diff_in_diff,
                diff_in_diff_pct=diff_in_diff_pct,
            ))
        
        # Total participants = nunique captains (test + control) in analysis date range
        total_participants = None
        if "captain_id" in col_types:
            tp_query = f"""
                SELECT COUNT(DISTINCT captain_id) as n
                FROM read_parquet('{parquet_path}')
                WHERE cohort IN ('{payload.test_cohort}', '{payload.control_cohort}')
                {date_filter}
            """
            try:
                tp_row = con.execute(tp_query).fetchone()
                total_participants = int(tp_row[0]) if tp_row and tp_row[0] is not None else None
            except Exception:
                pass
        
        return InsightsResponse(
            time_series=time_series,
            summary=summary,
            total_participants=total_participants,
        )
    
    finally:
        con.close()


def compute_insights_pandas(payload: InsightsRequest, df: pd.DataFrame) -> InsightsResponse:
    """
    Compute insights using pandas (original implementation for small files).
    """
    if not payload.metrics:
        raise HTTPException(status_code=400, detail="No metrics selected")

    working = df.copy()

    # Build a unified frame with explicit cohort_type
    if payload.test_cohort and payload.control_cohort:
        test_df = get_cohort(working, payload.test_cohort).copy()
        control_df = get_cohort(working, payload.control_cohort).copy()
        test_df["cohort_type"] = "test"
        control_df["cohort_type"] = "control"
        working = pd.concat([test_df, control_df], ignore_index=True)
    else:
        raise HTTPException(status_code=400, detail="Both test_cohort and control_cohort are required")

    # Ensure date is datetime
    working["date"] = pd.to_datetime(working["date"], errors="coerce")
    if working["date"].isna().any():
        raise HTTPException(status_code=400, detail="Invalid date values found in session data")

    def _agg_series(s: pd.Series, agg: str) -> float:
        if agg == "count":
            return float(s.count())
        if agg == "nunique":
            return float(s.nunique(dropna=True))
        s_num = pd.to_numeric(s, errors="coerce")
        if agg == "sum":
            return float(s_num.sum(skipna=True))
        if agg == "mean":
            return float(s_num.mean(skipna=True))
        if agg == "median":
            return float(s_num.median(skipna=True))
        raise ValueError(f"Unsupported aggregation: {agg}")

    # -------- Timeseries window (union of selected date ranges, if provided) --------
    start_candidates = []
    end_candidates = []
    for period in [payload.pre_period, payload.post_period]:
        if period and period.start_date:
            start_candidates.append(pd.to_datetime(period.start_date))
        if period and period.end_date:
            end_candidates.append(pd.to_datetime(period.end_date))
    ts_start = min(start_candidates) if start_candidates else None
    ts_end = max(end_candidates) if end_candidates else None

    ts_df = filter_by_date_range(working, ts_start, ts_end, date_col="date") if (ts_start or ts_end) else working
    ts_df = ts_df.sort_values(["cohort_type", "date"]).reset_index(drop=True)

    ts_frames = []
    for spec in payload.metrics:
        col = spec.column
        agg = spec.agg_func
        if col not in ts_df.columns:
            continue

        grouped = ts_df.groupby(["cohort_type", "date"], dropna=False)[col].apply(lambda x: _agg_series(x, agg))
        out = grouped.reset_index().rename(columns={col: "value"})
        out["metric"] = col
        out["agg_func"] = agg
        ts_frames.append(out)

    if ts_frames:
        ts_out = pd.concat(ts_frames, ignore_index=True)
    else:
        ts_out = pd.DataFrame(columns=["date", "cohort_type", "metric", "agg_func", "value"])

    time_series = [
        InsightsTimeSeriesPoint(
            date=pd.to_datetime(r["date"]).strftime("%Y-%m-%d"),
            cohort_type=str(r["cohort_type"]),
            metric=str(r["metric"]),
            agg_func=str(r["agg_func"]),
            value=float(r["value"]) if pd.notna(r["value"]) else 0.0,
        )
        for r in ts_out.to_dict("records")
    ]

    # -------- Executive summary (pre vs post aggregates per cohort_type) --------
    pre_df = working
    post_df = working
    if payload.pre_period:
        pre_df = filter_by_date_range(pre_df, payload.pre_period.start_date, payload.pre_period.end_date, date_col="date")
    if payload.post_period:
        post_df = filter_by_date_range(post_df, payload.post_period.start_date, payload.post_period.end_date, date_col="date")

    summary: list[InsightsSummaryRow] = []
    for spec in payload.metrics:
        col = spec.column
        agg = spec.agg_func
        if col not in working.columns:
            continue

        def _period_val(d: pd.DataFrame, cohort_type: str) -> float:
            subset = d.loc[d["cohort_type"] == cohort_type, col]
            return _agg_series(subset, agg) if not subset.empty else 0.0

        control_pre = _period_val(pre_df, "control")
        control_post = _period_val(post_df, "control")
        test_pre = _period_val(pre_df, "test")
        test_post = _period_val(post_df, "test")

        control_delta = control_post - control_pre
        test_delta = test_post - test_pre
        diff_in_diff = test_delta - control_delta

        control_delta_pct = (control_delta / control_pre * 100.0) if control_pre != 0 else None
        test_delta_pct = (test_delta / test_pre * 100.0) if test_pre != 0 else None
        diff_in_diff_pct = (diff_in_diff / control_pre * 100.0) if control_pre != 0 else None

        summary.append(
            InsightsSummaryRow(
                metric=col,
                agg_func=agg,
                control_pre=control_pre,
                control_post=control_post,
                control_delta=control_delta,
                control_delta_pct=control_delta_pct,
                test_pre=test_pre,
                test_post=test_post,
                test_delta=test_delta,
                test_delta_pct=test_delta_pct,
                diff_in_diff=diff_in_diff,
                diff_in_diff_pct=diff_in_diff_pct,
            )
        )

    return InsightsResponse(time_series=time_series, summary=summary)


@app.post("/insights", response_model=InsightsResponse, responses={400: {"model": ErrorResponse}})
def compute_insights(payload: InsightsRequest, x_session_id: Optional[str] = Header(default=None)) -> InsightsResponse:
    """
    Multi-metric Insights endpoint.
    - Returns daily timeseries for selected metrics (each with its own aggregation).
    - Returns an executive summary for pre vs post across test/control + deltas.
    
    Automatically uses DuckDB for large files (memory efficient) or pandas for small files.
    """
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    
    session = SESSION_STORE[x_session_id]
    
    # Check if it's a DuckDB session
    if isinstance(session, dict) and "parquet_path" in session:
        # Use DuckDB for memory-efficient query execution
        return compute_insights_duckdb(payload, session["parquet_path"])
    
    # Otherwise use pandas (for small files stored as DataFrames)
    if isinstance(session, pd.DataFrame):
        return compute_insights_pandas(payload, session)
    
    raise HTTPException(status_code=400, detail="Invalid session data format")


@app.post("/metrics", response_model=MetricsResponse, responses={400: {"model": ErrorResponse}})
def compute_metrics(payload: MetricsRequest, df: pd.DataFrame = Depends(get_session_df)) -> MetricsResponse:
    working = df.copy()

    # Apply cohort subset if provided
    cohorts = [c for c in [payload.test_cohort, payload.control_cohort] if c]
    if cohorts:
        working = subset_by_cohorts(working, cohorts)

    # Pre and Post period filters are used to compute summary aggregations per period
    pre_df = working
    post_df = working
    if payload.pre_period:
        pre_df = filter_by_date_range(working, payload.pre_period.start_date, payload.pre_period.end_date)
    if payload.post_period:
        post_df = filter_by_date_range(working, payload.post_period.start_date, payload.post_period.end_date)

    # Time series for plotting (full range filtered only by cohorts)
    ts_df = working.copy()
    ts_df = ts_df.sort_values(["cohort", "date"]).reset_index(drop=True)

    # Rolling windows
    for w in payload.rolling_windows:
        ts_df = rolling_average(ts_df, value_col="metric_value", window=w, by=["cohort"], date_col="date")

    # Normalized growth relative to baseline date or first date
    ts_df = normalized_growth(
        ts_df,
        value_col="metric_value",
        baseline_date=payload.normalized_growth_baseline_date,
        by=["cohort"],
        date_col="date",
    )

    # Build timeseries response points
    def _row_to_point(row) -> TimeSeriesPoint:
        point = {
            "date": pd.to_datetime(row["date"]).strftime("%Y-%m-%d"),
            "cohort": row["cohort"],
            "metric_value": float(row["metric_value"]),
            "metric_value_pct_change": float(row.get("metric_value_pct_change", float("nan"))) if "metric_value_pct_change" in row else None,
        }
        for w in payload.rolling_windows:
            col = f"metric_value_roll_{w}"
            if col in ts_df.columns:
                val = row.get(col)
                point[col] = float(val) if pd.notna(val) else None
        return TimeSeriesPoint(**point)

    time_series = [_row_to_point(r) for r in ts_df.to_dict("records")]

    # Compute summaries for pre and post periods by cohort
    summaries = []
    def _compute_agg(d: pd.DataFrame, label: str):
        for agg in payload.aggregations:
            out = aggregate_time_series(d, group_by=["cohort"], value_col="metric_value", agg=agg)
            # Map to test/control values
            test_val = out.loc[out["cohort"] == payload.test_cohort, f"metric_value_{agg}"]
            control_val = out.loc[out["cohort"] == payload.control_cohort, f"metric_value_{agg}"]
            tv = float(test_val.iloc[0]) if not test_val.empty else 0.0
            cv = float(control_val.iloc[0]) if not control_val.empty else 0.0
            mean_diff = tv - cv
            pct_change = (mean_diff / cv * 100.0) if cv != 0 else None
            summaries.append({
                "aggregation": agg,
                "test_value": tv,
                "control_value": cv,
                "mean_difference": mean_diff,
                "pct_change": pct_change,
            })

    _compute_agg(pre_df, "pre")
    _compute_agg(post_df, "post")

    return MetricsResponse(
        time_series=time_series,
        summaries=summaries,
    )


@app.post("/funnel", response_model=FunnelResponse, responses={400: {"model": ErrorResponse}})
def funnel(payload: FunnelRequest, df: pd.DataFrame = Depends(get_session_df)) -> FunnelResponse:
    working = df.copy()
    
    # Apply confirmation filtering if specified (support per-test/control and legacy)
    test_confirmed = getattr(payload, 'test_confirmed', None) or getattr(payload, 'confirmed', None) or ''
    control_confirmed = getattr(payload, 'control_confirmed', None) or getattr(payload, 'confirmed', None) or ''
    
    # Filter for test and control cohorts with optional confirmation filtering
    if payload.test_cohort and payload.control_cohort:
        test_data = get_cohort(working, payload.test_cohort, test_confirmed).copy()
        control_data = get_cohort(working, payload.control_cohort, control_confirmed).copy()
        # Label cohorts explicitly so identical cohort names remain distinct
        test_data["cohort"] = test_data["cohort"].astype(str).apply(lambda c: f"TEST: {c}")
        control_data["cohort"] = control_data["cohort"].astype(str).apply(lambda c: f"CONTROL: {c}")
        working = pd.concat([test_data, control_data], ignore_index=True)
    elif payload.test_cohort:
        working = get_cohort(working, payload.test_cohort, test_confirmed).copy()
        working["cohort"] = working["cohort"].astype(str).apply(lambda c: f"TEST: {c}")
    elif payload.control_cohort:
        working = get_cohort(working, payload.control_cohort, control_confirmed).copy()
        working["cohort"] = working["cohort"].astype(str).apply(lambda c: f"CONTROL: {c}")
    else:
        # If no specific cohorts, just apply confirmation filter if specified
        if test_confirmed:
            if test_confirmed not in working.columns:
                raise HTTPException(status_code=400, detail=f"Confirmation column '{test_confirmed}' not found")
            working = working[~working[test_confirmed].isna()]

    # Handle series breakout if specified
    series_breakout_col = payload.series_breakout
    if series_breakout_col:
        if series_breakout_col not in working.columns:
            raise HTTPException(status_code=400, detail=f"Series breakout column '{series_breakout_col}' not found in dataset. Available columns: {list(working.columns)[:20]}")
        # Check if column has any non-null values
        if working[series_breakout_col].isna().all():
            raise HTTPException(status_code=400, detail=f"Series breakout column '{series_breakout_col}' has no valid (non-null) values")
        # Store the series breakout column before computing timeseries
        series_values = working[series_breakout_col].astype(str).unique().tolist()
    else:
        series_values = None

    ts = compute_cohort_funnel_timeseries(working)
    metrics_available = [c for c in ts.columns if c not in {"date", "cohort"}]
    if not metrics_available:
        raise HTTPException(status_code=400, detail="No metrics available in dataset")

    # Determine the metric to use
    metric = payload.metric or metrics_available[0]
    agg = payload.agg or "sum"
    
    # If series breakout is specified, we need to recompute the metric with groupby
    # This is because precomputed metrics don't include the series breakout column
    if series_breakout_col:
        try:
            # Check if metric exists in working dataframe as a column
            if metric in working.columns:
                # Metric is a direct column, compute with groupby
                ts_with_series = compute_metric_timeseries_by_cohort(working, metric, agg, group_by=["cohort", series_breakout_col])
            else:
                # Metric is a precomputed one (like ao_days, online_days, etc.)
                # We need to compute it per series_breakout group from the original working data
                # by calling compute_cohort_funnel_timeseries on each group
                
                # Ensure date column exists in working
                working_for_series = working.copy()
                if "date" not in working_for_series.columns:
                    if "time" in working_for_series.columns:
                        working_for_series["date"] = pd.to_datetime(working_for_series["time"].astype(str), format="%Y%m%d", errors="coerce")
                    else:
                        raise HTTPException(status_code=400, detail="Cannot compute series breakout: missing 'date' or 'time' column")
                
                # Group working by series_breakout and compute timeseries for each group
                frames = []
                series_vals = working_for_series[series_breakout_col].dropna().unique()
                if len(series_vals) == 0:
                    raise HTTPException(status_code=400, detail=f"Series breakout column '{series_breakout_col}' has no valid values")
                
                for series_val in series_vals:
                    group_df = working_for_series[working_for_series[series_breakout_col] == series_val].copy()
                    if len(group_df) == 0:
                        continue
                    try:
                        group_ts = compute_cohort_funnel_timeseries(group_df)
                        if metric in group_ts.columns:
                            group_ts[series_breakout_col] = str(series_val)
                            frames.append(group_ts[["date", "cohort", metric, series_breakout_col]])
                    except Exception as group_error:
                        # Log but continue with other groups
                        print(f"Warning: Failed to compute timeseries for series value '{series_val}': {group_error}")
                        continue
                
                if frames:
                    ts_with_series = pd.concat(frames, ignore_index=True)
                else:
                    raise HTTPException(status_code=400, detail=f"Cannot compute series breakout for metric '{metric}': metric not found in any grouped data")
            
            # Replace ts with ts_with_series (which includes series_breakout)
            ts = ts_with_series.copy()
            
        except HTTPException:
            raise
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=400, detail=f"Failed to compute series breakout: {str(e)}")
    elif payload.metric and payload.metric not in metrics_available:
        # No series breakout, but metric needs to be computed
        try:
            extra = compute_metric_timeseries_by_cohort(working, payload.metric, agg)
        except Exception as e:
            raise HTTPException(status_code=400, detail=str(e))
        # Merge with existing to reuse date filtering & response shape
        ts = ts.copy()
        merged = ts.merge(extra, on=["date", "cohort"], how="outer", suffixes=("", "_extra"))
        # If both existed, prefer the non-null new computation
        if f"{metric}_extra" in merged.columns and metric in merged.columns:
            merged[metric] = merged[metric].fillna(merged[f"{metric}_extra"]).astype(float)
            merged = merged.drop(columns=[f"{metric}_extra"]) 
        elif f"{metric}_extra" in merged.columns and metric not in merged.columns:
            merged = merged.rename(columns={f"{metric}_extra": metric})
        ts = merged
        metrics_available = [c for c in ts.columns if c not in {"date", "cohort"}]

    pre_df = ts
    post_df = ts
    if payload.pre_period:
        pre_df = filter_by_date_range(ts, payload.pre_period.start_date, payload.pre_period.end_date, date_col="date")
    if payload.post_period:
        post_df = filter_by_date_range(ts, payload.post_period.start_date, payload.post_period.end_date, date_col="date")

    def to_points(d: pd.DataFrame) -> list[FunnelPoint]:
        points = []
        for r in d.to_dict("records"):
            series_val = str(r.get(series_breakout_col, "")) if series_breakout_col else None
            points.append(FunnelPoint(
                date=pd.to_datetime(r["date"]).strftime("%Y-%m-%d"),
                cohort=r["cohort"],
                metric=metric,
                value=float(r.get(metric, 0.0)),
                series_value=series_val
            ))
        return points

    def summarize(d: pd.DataFrame) -> dict[str, float]:
        if d.empty:
            return {}
        tmp = d.groupby("cohort")[metric].sum().reset_index()
        return {str(row["cohort"]): float(row[metric]) for _, row in tmp.iterrows()}

    return FunnelResponse(
        metrics_available=metrics_available,
        pre_series=to_points(pre_df),
        post_series=to_points(post_df),
        pre_summary=summarize(pre_df),
        post_summary=summarize(post_df),
    )


@app.get("/cohort-aggregation", response_model=CohortAggregationResponse, responses={400: {"model": ErrorResponse}})
def get_cohort_aggregation(df: pd.DataFrame = Depends(get_session_df)) -> CohortAggregationResponse:
    """Get cohort-level aggregation data as shown in the example"""
    working = df.copy()

    # Check if required columns exist
    required_columns = [
        "totalExpCaps", "visitedCaps", "clickedCaptain", "count_captain_pitch_centre_card_clicked_city","count_captain_pitch_centre_card_visible_city", "exploredCaptains",
        "exploredCaptains_Subs", "exploredCaptains_EPKM", "exploredCaptains_FlatCommission",
        "exploredCaptains_CM", "confirmedCaptains", "confirmedCaptains_Subs",
        "confirmedCaptains_Subs_purchased", "confirmedCaptains_Subs_purchased_weekend",
        "confirmedCaptains_EPKM", "confirmedCaptains_FlatCommission", "confirmedCaptains_CM"
    ]

    missing_columns = [col for col in required_columns if col not in working.columns]
    if missing_columns:
        raise HTTPException(
            status_code=400,
            detail=f"Missing required columns for cohort aggregation: {missing_columns}. "
                   f"Available columns: {list(working.columns)}"
        )

    # Perform the aggregation as specified in the example
    result = working.groupby(["cohort"]).agg({
        "totalExpCaps": "nunique",
        "visitedCaps": "nunique",
        'clickedCaptain': 'nunique',
        'count_captain_pitch_centre_card_clicked_city': 'sum',
        'count_captain_pitch_centre_card_visible_city': 'sum',
        'exploredCaptains': 'nunique',
        'exploredCaptains_Subs': 'nunique',
        'exploredCaptains_EPKM': 'nunique',
        'exploredCaptains_FlatCommission': 'nunique',
        'exploredCaptains_CM': 'nunique',
        'confirmedCaptains': 'nunique',
        'confirmedCaptains_Subs': 'nunique',
        'confirmedCaptains_Subs_purchased': 'nunique',
        'confirmedCaptains_Subs_purchased_weekend': 'nunique',
        'confirmedCaptains_EPKM': 'nunique',
        'confirmedCaptains_FlatCommission': 'nunique',
        'confirmedCaptains_CM': 'nunique'

    }).reset_index().sort_values("exploredCaptains", ascending=False)

    # Calculate the ratio columns
    result['Visit2Click'] = result['clickedCaptain'] / result['visitedCaps']
    result['Base2Visit'] = result['visitedCaps'] / result['totalExpCaps']
    # To fix division by zero, use np.where to safely compute Click2Confirm
    import numpy as np
    result['Click2Confirm'] = np.where(result['clickedCaptain'] == 0, 0, result['confirmedCaptains'] / result['clickedCaptain'])

    # Handle division by zero
    result['Visit2Click'] = result['Visit2Click'].fillna(0)
    result['Base2Visit'] = result['Base2Visit'].fillna(0)
    result['Click2Confirm'] = result['Click2Confirm'].fillna(0)
    # Convert to list of CohortAggregationRow objects
    data = []
    for _, row in result.iterrows():
        data.append(CohortAggregationRow(
            cohort=str(row['cohort']),
            totalExpCaps=float(row['totalExpCaps']),
            visitedCaps=float(row['visitedCaps']),
            clickedCaptain=float(row['clickedCaptain']),
            pitch_centre_card_clicked=float(row['count_captain_pitch_centre_card_clicked_city']),
            pitch_centre_card_visible=float(row['count_captain_pitch_centre_card_visible_city']),
            exploredCaptains=float(row['exploredCaptains']),
            exploredCaptains_Subs=float(row['exploredCaptains_Subs']),
            exploredCaptains_EPKM=float(row['exploredCaptains_EPKM']),
            exploredCaptains_FlatCommission=float(row['exploredCaptains_FlatCommission']),
            exploredCaptains_CM=float(row['exploredCaptains_CM']),
            confirmedCaptains=float(row['confirmedCaptains']),
            confirmedCaptains_Subs=float(row['confirmedCaptains_Subs']),
            confirmedCaptains_Subs_purchased=float(row['confirmedCaptains_Subs_purchased']),
            confirmedCaptains_Subs_purchased_weekend=float(row['confirmedCaptains_Subs_purchased_weekend']),
            confirmedCaptains_EPKM=float(row['confirmedCaptains_EPKM']),
            confirmedCaptains_FlatCommission=float(row['confirmedCaptains_FlatCommission']),
            confirmedCaptains_CM=float(row['confirmedCaptains_CM']),
            Visit2Click=float(row['Visit2Click']),
            Base2Visit=float(row['Base2Visit']),
            Click2Confirm=float(row['Click2Confirm']),
        ))

    return CohortAggregationResponse(data=data)


@app.delete("/session")
def clear_session(x_session_id: Optional[str] = Header(default=None)):
    if x_session_id:
        # Check if it's a DuckDB session and clean up parquet file
        if x_session_id in SESSION_STORE:
            session = SESSION_STORE[x_session_id]
            if isinstance(session, dict) and "parquet_path" in session:
                # DuckDB session - clean up parquet file
                parquet_path = session.get("parquet_path")
                if parquet_path and os.path.exists(parquet_path):
                    try:
                        os.unlink(parquet_path)
                    except Exception:
                        pass
                # Also clean up the temp directory
                temp_dir = os.path.dirname(parquet_path) if parquet_path else None
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
            del SESSION_STORE[x_session_id]
        
        # Clean up temp file if exists (for pandas sessions)
        if x_session_id in SESSION_FILE_STORE:
            temp_path = SESSION_FILE_STORE[x_session_id]
            # For DuckDB sessions, temp_path might be the directory
            if os.path.isdir(temp_path):
                shutil.rmtree(temp_path, ignore_errors=True)
            else:
                temp_dir = os.path.dirname(temp_path)
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
            del SESSION_FILE_STORE[x_session_id]
    
    return {"ok": True}


@app.post("/statistical-test")
def run_statistical_test_endpoint(
    request: StatTestRequest,
    x_session_id: Optional[str] = Header(default=None)
) -> StatTestResult:
    """Run statistical tests on cohort data"""
    from statistical_analysis import run_statistical_test
    
    try:
        result = run_statistical_test(request)
        return result
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.post("/captain-level-aggregation", response_model=CaptainLevelResponse, responses={400: {"model": ErrorResponse}})
def captain_level_aggregation(
    payload: CaptainLevelRequest,
    df: pd.DataFrame = Depends(get_session_df)
) -> CaptainLevelResponse:
    """
    Captain-level aggregation grouped by time and a categorical column.
    Filters data by cohort and confirmation level, then aggregates metrics.
    """
    working = df.copy()
    
    # Ensure date column is present
    if "date" not in working.columns:
        raise HTTPException(status_code=400, detail="Dataset must have 'date' column")
    
    # Ensure group_by_column exists
    if payload.group_by_column not in working.columns:
        raise HTTPException(
            status_code=400,
            detail=f"Group by column '{payload.group_by_column}' not found in dataset"
        )
    
    # Validate all metric columns exist
    for metric_agg in payload.metric_aggregations:
        if metric_agg.column not in working.columns:
            raise HTTPException(
                status_code=400,
                detail=f"Metric column '{metric_agg.column}' not found in dataset"
            )
    
    def filter_captain_level(cohort_name: str, confirmation_filter: Optional[str]) -> pd.DataFrame:
        """Filter data by cohort and optional confirmation level"""
        cohort_df = working[working["cohort"] == cohort_name].copy()
        if confirmation_filter and confirmation_filter in cohort_df.columns:
            cohort_df = cohort_df[~cohort_df[confirmation_filter].isna()]
        return cohort_df
    
    # Get test and control data
    test_df = filter_captain_level(payload.test_cohort, payload.test_confirmed)
    control_df = filter_captain_level(payload.control_cohort, payload.control_confirmed)
    
    if test_df.empty:
        raise HTTPException(status_code=400, detail=f"No data found for test cohort '{payload.test_cohort}'")
    if control_df.empty:
        raise HTTPException(status_code=400, detail=f"No data found for control cohort '{payload.control_cohort}'")
    
    # Helper to aggregate data
    def aggregate_data(data: pd.DataFrame, period: str, cohort_type: str) -> list[CaptainLevelAggregationRow]:
        """Group by date and group_by_column, then aggregate metrics"""
        if data.empty:
            return []
        
        # Build aggregation dict - map column names to aggregation functions
        agg_dict = {}
        agg_key_mapping = {}  # Maps original column to our custom key
        
        for metric_agg in payload.metric_aggregations:
            col_name = metric_agg.column
            agg_func = metric_agg.agg_func
            agg_key = f"{col_name}_{agg_func}"
            
            # For pandas groupby, we need column name as key
            if col_name not in agg_dict:
                agg_dict[col_name] = []
            agg_dict[col_name].append(agg_func)
            
            # Store mapping for later renaming
            agg_key_mapping[f"{col_name}_{agg_func}"] = agg_key
        
        # Group by date and the categorical column
        grouped = data.groupby(["date", payload.group_by_column]).agg(agg_dict).reset_index()
        
        # Flatten multi-level column names if they exist
        if isinstance(grouped.columns, pd.MultiIndex):
            new_cols = ["date", payload.group_by_column]
            for col in grouped.columns[2:]:  # Skip date and group_by_column
                if col[1]:  # If there's an aggregation function
                    new_cols.append(f"{col[0]}_{col[1]}")
                else:
                    new_cols.append(col[0])
            grouped.columns = new_cols
        
        # Convert to response format
        rows = []
        for _, row in grouped.iterrows():
            aggregations = {}
            for orig_key, custom_key in agg_key_mapping.items():
                val = row.get(orig_key)
                aggregations[custom_key] = float(val) if pd.notna(val) else 0.0
            
            rows.append(CaptainLevelAggregationRow(
                period=period,
                cohort_type=cohort_type,
                date=pd.to_datetime(row["date"]).strftime("%Y-%m-%d"),
                group_value=str(row[payload.group_by_column]),
                aggregations=aggregations
            ))
        
        return rows
    
    # Filter by date ranges
    pre_test = test_df
    post_test = test_df
    pre_control = control_df
    post_control = control_df
    
    if payload.pre_period:
        pre_test = filter_by_date_range(
            test_df,
            payload.pre_period.start_date,
            payload.pre_period.end_date,
            date_col="date"
        )
        pre_control = filter_by_date_range(
            control_df,
            payload.pre_period.start_date,
            payload.pre_period.end_date,
            date_col="date"
        )
    
    if payload.post_period:
        post_test = filter_by_date_range(
            test_df,
            payload.post_period.start_date,
            payload.post_period.end_date,
            date_col="date"
        )
        post_control = filter_by_date_range(
            control_df,
            payload.post_period.start_date,
            payload.post_period.end_date,
            date_col="date"
        )
    
    # Aggregate all combinations
    result_data = []
    result_data.extend(aggregate_data(pre_test, "pre", "test"))
    result_data.extend(aggregate_data(post_test, "post", "test"))
    result_data.extend(aggregate_data(pre_control, "pre", "control"))
    result_data.extend(aggregate_data(post_control, "post", "control"))
    
    # Extract metric names
    metrics = [f"{m.column}_{m.agg_func}" for m in payload.metric_aggregations]
    
    return CaptainLevelResponse(
        data=result_data,
        group_by_column=payload.group_by_column,
        metrics=metrics
    )


# ============================================================================
# FUNNEL ANALYSIS ENDPOINTS
# ============================================================================

@app.post("/funnel-analysis/upload-mobile-numbers", response_model=MobileNumberUploadResponse, responses={400: {"model": ErrorResponse}})
async def upload_mobile_numbers(file: UploadFile = File(...)) -> MobileNumberUploadResponse:
    """
    Upload CSV with mobile_number column (and optional cohort column) for funnel analysis
    """
    if not file.filename.lower().endswith(".csv"):
        raise HTTPException(status_code=400, detail="Only CSV files are supported")
    
    content = await file.read()
    try:
        df = pd.read_csv(io.BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to read CSV: {exc}")
    
    # Validate that mobile_number column exists
    if "mobile_number" not in df.columns:
        raise HTTPException(status_code=400, detail="CSV must include 'mobile_number' column")
    
    # Check if cohort column exists
    has_cohort = "cohort" in df.columns
    
    # Ensure mobile_number is treated as string
    df['mobile_number'] = df['mobile_number'].astype(str)
    
    if has_cohort:
        df['cohort'] = df['cohort'].astype(str)
    
    # Drop duplicates based on mobile_number (and cohort if present)
    original_rows = len(df)
    if has_cohort:
        df = df.drop_duplicates(subset=['mobile_number', 'cohort'], keep='first')
    else:
        df = df.drop_duplicates(subset=['mobile_number'], keep='first')
    
    duplicates_removed = original_rows - len(df)
    
    # Generate session ID and store
    funnel_session_id = secrets.token_hex(16)
    FUNNEL_SESSION_STORE[funnel_session_id] = df
    
    # Get preview (first 5 rows)
    preview = df.head(5).to_dict('records')
    
    return MobileNumberUploadResponse(
        funnel_session_id=funnel_session_id,
        num_rows=len(df),
        columns=list(df.columns),
        has_cohort=has_cohort,
        preview=preview,
        duplicates_removed=duplicates_removed
    )


@app.post("/funnel-analysis/get-captain-ids", response_model=CaptainIdResponse, responses={400: {"model": ErrorResponse}})
async def get_captain_ids(
    payload: CaptainIdRequest,
    x_funnel_session_id: Optional[str] = Header(default=None)
) -> CaptainIdResponse:
    """
    Fetch captain IDs for mobile numbers in current session
    """
    if not x_funnel_session_id or x_funnel_session_id not in FUNNEL_SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing funnel_session_id. Upload mobile numbers first.")
    
    mobile_number_df = FUNNEL_SESSION_STORE[x_funnel_session_id]
    
    try:
        from funnel import get_captain_id
        result_df = get_captain_id(mobile_number_df, payload.username)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch captain IDs from Presto: {exc}")
    
    # Update session with new dataframe including captain_id
    FUNNEL_SESSION_STORE[x_funnel_session_id] = result_df
    
    # Count how many captain IDs were found (non-null)
    num_captains_found = len(result_df['captain_id'].dropna().unique())
    
    # Get preview
    preview = result_df.head(5).to_dict('records')
    
    return CaptainIdResponse(
        num_rows=len(result_df),
        num_captains_found=int(num_captains_found),
        preview=preview
    )


@app.post("/funnel-analysis/get-ao-funnel", response_model=AOFunnelResponse, responses={400: {"model": ErrorResponse}})
async def get_ao_funnel_data(
    payload: AOFunnelRequest,
    x_funnel_session_id: Optional[str] = Header(default=None)
) -> AOFunnelResponse:
    """
    Fetch AO funnel metrics for captain IDs in current session
    """
    if not x_funnel_session_id or x_funnel_session_id not in FUNNEL_SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing funnel_session_id. Get captain IDs first.")
    
    captain_id_df = FUNNEL_SESSION_STORE[x_funnel_session_id]
    
    # Validate that captain_id column exists
    if 'captain_id' not in captain_id_df.columns:
        raise HTTPException(status_code=400, detail="No captain_id column found. Run 'Get Captain IDs' first.")
    
    # Check if there are any valid captain IDs
    if captain_id_df['captain_id'].isna().all():
        raise HTTPException(status_code=400, detail="No valid captain IDs found. Cannot fetch funnel data.")
    
    try:
        from funnel import get_ao_funnel
        result_df = get_ao_funnel(
            captain_id_df,
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.time_level,
            payload.tod_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch AO funnel data from Presto: {exc}")
    
    # Format the dataframe for cohort analysis
    # Ensure we have a proper date column
    if 'time' in result_df.columns and 'date' not in result_df.columns:
        # Convert time column (YYYYMMDD format) to date
        result_df['date'] = pd.to_datetime(result_df['time'], format='%Y%m%d', errors='coerce')
    elif 'time' in result_df.columns:
        # Also ensure date column is datetime
        result_df['date'] = pd.to_datetime(result_df['date'], errors='coerce')
    
    # Ensure cohort column exists, if not create from existing data
    if 'cohort' not in result_df.columns:
        # If no cohort column, create a default one
        result_df['cohort'] = 'all_captains'
    
    # Update session with funnel data
    FUNNEL_SESSION_STORE[x_funnel_session_id] = result_df
    
    # Identify metric columns (exclude identifier columns)
    exclude_cols = {'mobile_number', 'captain_id', 'cohort', 'city', 'time', 'date'}
    metric_cols = [c for c in result_df.columns if c not in exclude_cols]
    
    # Calculate unique captain IDs from full dataset
    unique_captain_ids = int(result_df['captain_id'].nunique())
    
    # Get preview (first 10 rows)
    preview = result_df.head(10).to_dict('records')
    
    return AOFunnelResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        preview=preview,
        metrics=metric_cols,
        unique_captain_ids=unique_captain_ids
    )


@app.post("/funnel-analysis/use-for-analysis", response_model=UploadResponse, responses={400: {"model": ErrorResponse}})
def use_funnel_for_analysis(x_funnel_session_id: Optional[str] = Header(default=None)) -> UploadResponse:
    """
    Transfer funnel analysis data to main cohort analysis session
    This allows using the AO funnel data for plotting and further analysis
    """
    if not x_funnel_session_id or x_funnel_session_id not in FUNNEL_SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing funnel_session_id. Complete funnel analysis first.")
    
    # Get the funnel dataframe
    funnel_df = FUNNEL_SESSION_STORE[x_funnel_session_id].copy()
    
    # Ensure date column is properly formatted
    if 'date' not in funnel_df.columns and 'time' in funnel_df.columns:
        funnel_df['date'] = pd.to_datetime(funnel_df['time'], format='%Y%m%d', errors='coerce')
    elif 'date' in funnel_df.columns:
        funnel_df['date'] = pd.to_datetime(funnel_df['date'], errors='coerce')
    else:
        raise HTTPException(status_code=400, detail="No date/time column found in funnel data")
    
    # Ensure cohort column exists
    if 'cohort' not in funnel_df.columns:
        funnel_df['cohort'] = 'all_captains'
    
    # Ensure cohort is string
    funnel_df['cohort'] = funnel_df['cohort'].astype(str)
    
    # Drop any rows with invalid dates
    invalid_dates = funnel_df['date'].isna().sum()
    if invalid_dates > 0:
        funnel_df = funnel_df[~funnel_df['date'].isna()]
    
    # Create a new session in the main store
    session_id = secrets.token_hex(16)
    SESSION_STORE[session_id] = funnel_df
    
    # Get cohorts and date range
    cohorts = sorted(funnel_df["cohort"].dropna().unique().tolist())
    date_min = funnel_df["date"].min()
    date_max = funnel_df["date"].max()
    
    # Determine available metric columns
    exclude_cols = {'cohort', 'date', 'time', 'mobile_number', 'captain_id', 'city'}
    metric_candidates = [c for c in funnel_df.columns if c not in exclude_cols]
    
    return UploadResponse(
        session_id=session_id,
        num_rows=funnel_df.shape[0],
        columns=list(funnel_df.columns.astype(str)),
        cohorts=cohorts,
        date_min=date_min.strftime("%Y-%m-%d"),
        date_max=date_max.strftime("%Y-%m-%d"),
        metrics=metric_candidates,
    )


@app.get("/funnel-analysis/export-csv")
def export_funnel_csv(x_funnel_session_id: Optional[str] = Header(default=None)):
    """Export full funnel dataset as CSV"""
    from fastapi.responses import StreamingResponse
    
    if not x_funnel_session_id or x_funnel_session_id not in FUNNEL_SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing funnel_session_id")
    
    df = FUNNEL_SESSION_STORE[x_funnel_session_id]
    
    # Convert to CSV
    csv_buffer = io.StringIO()
    df.to_csv(csv_buffer, index=False)
    csv_buffer.seek(0)
    
    return StreamingResponse(
        iter([csv_buffer.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=funnel_data.csv"}
    )


@app.post("/funnel-analysis/dapr-bucket", response_model=DaprBucketResponse, responses={400: {"model": ErrorResponse}})
async def get_dapr_bucket(payload: DaprBucketRequest) -> DaprBucketResponse:
    """
    Fetch DAPR bucket distribution data from Presto
    """
    try:
        from funnel import dapr_bucket
        result_df = dapr_bucket(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service_category,
            payload.low_dapr,
            payload.high_dapr
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch DAPR bucket data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return DaprBucketResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/fe2net", response_model=Fe2NetResponse, responses={400: {"model": ErrorResponse}})
async def get_fe2net(payload: Fe2NetRequest) -> Fe2NetResponse:
    """
    Fetch FE2Net funnel data from Presto
    """
    try:
        from funnel import fe2net
        result_df = fe2net(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service_category,
            payload.geo_level,
            payload.time_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch FE2Net data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return Fe2NetResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/rtu-performance", response_model=RtuPerformanceResponse, responses={400: {"model": ErrorResponse}})
async def get_rtu_performance(payload: RtuPerformanceRequest) -> RtuPerformanceResponse:
    """
    Fetch RTU Performance metrics from Presto
    """
    try:
        from funnel import performance_metrics
        result_df = performance_metrics(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.perf_cut,
            payload.consistency_cut,
            payload.time_level,
            payload.tod_level,
            payload.service_category
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch RTU Performance data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return RtuPerformanceResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/r2a", response_model=R2AResponse, responses={400: {"model": ErrorResponse}})
async def get_r2a(payload: R2ARequest) -> R2AResponse:
    """
    Fetch R2A% (Registration to Activation) metrics from Presto
    """
    try:
        from funnel import r2a_registration_by_activation
        result_df = r2a_registration_by_activation(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service,
            payload.time_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch R2A data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return R2AResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/r2a-percentage", response_model=R2APercentageResponse, responses={400: {"model": ErrorResponse}})
async def get_r2a_percentage(payload: R2APercentageRequest) -> R2APercentageResponse:
    """
    Fetch R2A% metrics from Presto
    """
    try:
        from funnel import r2a_pecentage
        result_df = r2a_pecentage(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service,
            payload.time_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch R2A% data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return R2APercentageResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.post("/captain-dashboards/a2phh-summary", response_model=A2PhhSummaryResponse, responses={400: {"model": ErrorResponse}})
async def get_a2phh_summary(payload: A2PhhSummaryRequest) -> A2PhhSummaryResponse:
    """
    Fetch A2PHH Summary M0 metrics from Presto
    """
    try:
        from funnel import a2phh_summary
        result_df = a2phh_summary(
            payload.username,
            payload.start_date,
            payload.end_date,
            payload.city,
            payload.service,
            payload.time_level
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to fetch A2PHH Summary data: {exc}")
    
    # Convert all data to records
    data = result_df.to_dict('records')
    
    return A2PhhSummaryResponse(
        num_rows=len(result_df),
        columns=list(result_df.columns),
        data=data
    )


@app.delete("/funnel-analysis/session")
def clear_funnel_session(x_funnel_session_id: Optional[str] = Header(default=None)):
    """Clear funnel analysis session"""
    if x_funnel_session_id and x_funnel_session_id in FUNNEL_SESSION_STORE:
        del FUNNEL_SESSION_STORE[x_funnel_session_id]
    return {"ok": True}


# ============================================================================
# REPORT BUILDER ENDPOINTS
# ============================================================================

@app.post("/report/create")
def create_report() -> dict:
    """Create a new report session"""
    report_id = secrets.token_hex(16)
    REPORT_STORE[report_id] = []
    return {"report_id": report_id}


@app.post("/report/add-item", response_model=ReportAddResponse)
def add_report_item(
    payload: ReportAddRequest,
    x_report_id: Optional[str] = Header(default=None, alias="x-report-id")
) -> ReportAddResponse:
    """Add an item (chart, table, or text) to the report"""
    from datetime import datetime
    
    print(f"DEBUG: add_report_item called with x_report_id={x_report_id}")
    print(f"DEBUG: Payload type={payload.type}, title={payload.title}")
    print(f"DEBUG: REPORT_STORE keys before: {list(REPORT_STORE.keys())}")
    
    # Create report if not exists
    if not x_report_id or x_report_id not in REPORT_STORE:
        if not x_report_id:
            x_report_id = secrets.token_hex(16)
        REPORT_STORE[x_report_id] = []
        print(f"DEBUG: Created new report with ID: {x_report_id}")
    
    item_id = secrets.token_hex(8)
    item = {
        "id": item_id,
        "type": payload.type,
        "title": payload.title,
        "content": payload.content,
        "comment": payload.comment or "",
        "timestamp": datetime.now().isoformat()
    }
    
    REPORT_STORE[x_report_id].append(item)
    print(f"DEBUG: Added item to report {x_report_id}, total items: {len(REPORT_STORE[x_report_id])}")
    print(f"DEBUG: REPORT_STORE keys after: {list(REPORT_STORE.keys())}")
    
    return ReportAddResponse(
        report_id=x_report_id,
        item_id=item_id,
        num_items=len(REPORT_STORE[x_report_id])
    )


@app.put("/report/update-comment")
def update_report_comment(
    payload: ReportUpdateCommentRequest,
    x_report_id: Optional[str] = Header(default=None, alias="x-report-id")
):
    """Update comment for a report item"""
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    for item in items:
        if item["id"] == payload.item_id:
            item["comment"] = payload.comment
            return {"ok": True}
    
    raise HTTPException(status_code=404, detail="Item not found in report")


@app.put("/report/update-title")
def update_report_title(
    payload: ReportUpdateTitleRequest,
    x_report_id: Optional[str] = Header(default=None, alias="x-report-id")
):
    """Update title for a report item"""
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    for item in items:
        if item["id"] == payload.item_id:
            item["title"] = payload.title
            return {"ok": True}
    
    raise HTTPException(status_code=404, detail="Item not found in report")


@app.delete("/report/item/{item_id}")
def delete_report_item(
    item_id: str,
    x_report_id: Optional[str] = Header(default=None, alias="x-report-id")
):
    """Delete an item from the report"""
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    REPORT_STORE[x_report_id] = [item for item in items if item["id"] != item_id]
    
    return {"ok": True, "num_items": len(REPORT_STORE[x_report_id])}


@app.get("/report/list", response_model=ReportListResponse)
def list_report_items(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")) -> ReportListResponse:
    """Get all items in the current report"""
    print(f"DEBUG: list_report_items called with x_report_id={x_report_id}")
    print(f"DEBUG: REPORT_STORE keys: {list(REPORT_STORE.keys())}")
    
    if not x_report_id:
        print("DEBUG: No report ID provided, returning empty")
        return ReportListResponse(report_id="", items=[])
    
    if x_report_id not in REPORT_STORE:
        print(f"DEBUG: Report ID {x_report_id} not found in store, returning empty")
        return ReportListResponse(report_id=x_report_id, items=[])
    
    items = [ReportItem(**item) for item in REPORT_STORE[x_report_id]]
    print(f"DEBUG: Returning {len(items)} items for report {x_report_id}")
    return ReportListResponse(report_id=x_report_id, items=items)


@app.get("/report/export", response_model=ReportExportResponse)
def export_report(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")) -> ReportExportResponse:
    """Export report as HTML document"""
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    
    # Build HTML document
    html_parts = [
        """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Experiment Report</title>
            <style>
                body {
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 40px 20px;
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                }
                .container {
                    background: white;
                    padding: 40px;
                    border-radius: 12px;
                    box-shadow: 0 20px 60px rgba(0,0,0,0.3);
                }
                h1 {
                    color: #1e293b;
                    font-size: 32px;
                    margin-bottom: 10px;
                    border-bottom: 3px solid #8b5cf6;
                    padding-bottom: 10px;
                }
                .subtitle {
                    color: #64748b;
                    font-size: 14px;
                    margin-bottom: 30px;
                }
                .report-item {
                    margin: 30px 0;
                    padding: 25px;
                    border: 2px solid #e2e8f0;
                    border-radius: 8px;
                    background: #f8fafc;
                }
                .item-header {
                    font-size: 20px;
                    font-weight: 700;
                    color: #334155;
                    margin-bottom: 10px;
                }
                .item-type {
                    display: inline-block;
                    padding: 4px 12px;
                    background: #8b5cf6;
                    color: white;
                    border-radius: 12px;
                    font-size: 12px;
                    font-weight: 600;
                    margin-bottom: 12px;
                }
                .item-comment {
                    margin-top: 15px;
                    padding: 15px;
                    background: white;
                    border-left: 4px solid #8b5cf6;
                    font-style: italic;
                    color: #475569;
                }
                .item-timestamp {
                    font-size: 12px;
                    color: #94a3b8;
                    margin-top: 8px;
                }
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin-top: 15px;
                    font-size: 13px;
                }
                th {
                    background: #8b5cf6;
                    color: white;
                    padding: 10px;
                    text-align: left;
                    font-weight: 600;
                }
                td {
                    padding: 10px;
                    border-bottom: 1px solid #e2e8f0;
                }
                tr:nth-child(even) {
                    background: #f1f5f9;
                }
                .chart-config {
                    background: #f1f5f9;
                    padding: 15px;
                    border-radius: 6px;
                    margin-top: 15px;
                    font-family: 'Courier New', monospace;
                    font-size: 12px;
                }
                .page-break {
                    page-break-after: always;
                }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>📊 Experiment Report</h1>
                <p class="subtitle">Generated on """ + pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S") + """</p>
        """
    ]
    
    # Add each report item
    for idx, item in enumerate(items):
        item_html = f"""
            <div class="report-item">
                <span class="item-type">{item['type'].upper()}</span>
                <div class="item-header">{idx + 1}. {item['title']}</div>
        """
        
        # Render content based on type
        if item['type'] == 'table':
            # Render table image if available, otherwise render table data
            content = item['content']
            if 'imageDataUrl' in content and content['imageDataUrl']:
                # Display the actual table image
                img_data = content['imageDataUrl']
                item_html += f"<div style='margin-top:15px; text-align:center;'><img src='{img_data}' alt='{item['title']}' style='max-width:100%; height:auto; border:1px solid #e2e8f0; border-radius:8px;' /></div>"
            elif 'data' in content and len(content['data']) > 0:
                # Fallback to table data if no image
                # Get columns from first row
                columns = list(content['data'][0].keys())
                item_html += "<table>"
                item_html += "<thead><tr>"
                for col in columns:
                    item_html += f"<th>{col}</th>"
                item_html += "</tr></thead><tbody>"
                
                # Add rows (limit to first 50 for readability)
                for row in content['data'][:50]:
                    item_html += "<tr>"
                    for col in columns:
                        val = row.get(col, '')
                        # Format numbers
                        if isinstance(val, (int, float)):
                            val = f"{val:,.2f}" if isinstance(val, float) else f"{val:,}"
                        item_html += f"<td>{val}</td>"
                    item_html += "</tr>"
                item_html += "</tbody></table>"
                
                if len(content['data']) > 50:
                    item_html += f"<p style='margin-top:10px; color:#64748b; font-size:12px;'>Showing first 50 of {len(content['data'])} rows</p>"
        
        elif item['type'] == 'chart':
            # Render chart image if available
            content = item['content']
            if 'imageDataUrl' in content and content['imageDataUrl']:
                # Display the actual chart image
                img_data = content['imageDataUrl']
                item_html += f"<div style='margin-top:15px; text-align:center;'><img src='{img_data}' alt='{item['title']}' style='max-width:100%; height:auto; border:1px solid #e2e8f0; border-radius:8px;' /></div>"
            else:
                # Fallback to metadata if no image
                item_html += "<div class='chart-config'>"
                item_html += f"<strong>Chart Type:</strong> {content.get('chartType', 'N/A')}<br>"
                item_html += f"<strong>X-Axis:</strong> {content.get('xAxis', 'N/A')}<br>"
                item_html += f"<strong>Y-Axes:</strong> {', '.join(content.get('yAxes', []))}<br>"
                if content.get('seriesBy'):
                    item_html += f"<strong>Series Breakout:</strong> {content.get('seriesBy')}<br>"
                item_html += f"<strong>Data Points:</strong> {len(content.get('data', []))}"
                item_html += "</div>"
        
        elif item['type'] == 'text':
            # Render text content
            text_content = item['content'].get('text', '')
            item_html += f"<div style='margin-top:15px; padding:15px; background:white; border-radius:6px;'>{text_content}</div>"
        
        # Add comment if present
        if item['comment']:
            item_html += f"<div class='item-comment'>💬 {item['comment']}</div>"
        
        # Add timestamp
        item_html += f"<div class='item-timestamp'>Added: {item['timestamp']}</div>"
        
        item_html += "</div>"
        
        # Add page break after every 2 items for printing
        if (idx + 1) % 2 == 0:
            item_html += "<div class='page-break'></div>"
        
        html_parts.append(item_html)
    
    html_parts.append("""
            </div>
        </body>
        </html>
    """)
    
    return ReportExportResponse(report_html="".join(html_parts))


@app.delete("/report/clear")
def clear_report(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")):
    """Clear all items from the report"""
    if x_report_id and x_report_id in REPORT_STORE:
        REPORT_STORE[x_report_id] = []
    return {"ok": True}


@app.get("/report/export/pdf")
def export_report_pdf(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")):
    """Export report as PDF"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
        from reportlab.lib.enums import TA_LEFT
        import base64
        import tempfile
        import os
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"PDF export library not available: {str(e)}")
    
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    try:
        items = REPORT_STORE[x_report_id]
        temp_files = []  # Track temporary files for cleanup
        
        # Create temporary PDF file
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Container for the 'Flowable' objects
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='black',
            spaceAfter=12,
        )
        
        # Add title
        story.append(Paragraph("Experiment Report", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Add each item
        for idx, item in enumerate(items):
            # Item title
            story.append(Paragraph(f"{idx + 1}. {item['title']}", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            
            # Add image if present
            if item['type'] in ['chart', 'table'] and 'imageDataUrl' in item.get('content', {}):
                try:
                    # Decode base64 image
                    img_data = item['content']['imageDataUrl']
                    if img_data.startswith('data:image'):
                        img_data = img_data.split(',')[1]
                    
                    img_bytes = base64.b64decode(img_data)
                    
                    # Use PIL to open the image, then save to a temporary file that stays open
                    from PIL import Image as PILImage
                    pil_img = PILImage.open(io.BytesIO(img_bytes))
                    
                    # Save to temporary file that ReportLab can read
                    img_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    pil_img.save(img_file.name, format='PNG')
                    img_file.close()
                    
                    # Add image to PDF - ReportLab will read the file
                    img = RLImage(img_file.name, width=5*inch, height=3.75*inch)
                    story.append(img)
                    story.append(Spacer(1, 0.2*inch))
                    
                    # Clean up after PDF is built (will be done later)
                    # Store file path for cleanup
                    temp_files.append(img_file.name)
                except Exception as e:
                    print(f"Error adding image to PDF: {e}")
                    import traceback
                    traceback.print_exc()
            
            # Add text content
            if item['type'] == 'text' and 'text' in item.get('content', {}):
                story.append(Paragraph(item['content']['text'], styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            
            # Add comment if present
            if item.get('comment'):
                comment_style = ParagraphStyle(
                    'Comment',
                    parent=styles['Normal'],
                    fontStyle='italic',
                    textColor='#666666',
                    leftIndent=20,
                    backColor='#FFFACD',
                )
                story.append(Paragraph(f"Comment: {item['comment']}", comment_style))
            
            story.append(Spacer(1, 0.3*inch))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        # Clean up temporary image files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                print(f"Error cleaning up temp file {temp_file}: {e}")
        
        return Response(
            content=buffer.read(),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=experiment_report_{pd.Timestamp.now().strftime('%Y%m%d')}.pdf",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "x-report-id",
            }
        )
    except Exception as e:
        # Clean up temporary files on error
        if 'temp_files' in locals():
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.unlink(temp_file)
                except:
                    pass
        
        print(f"Error generating PDF: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")


@app.get("/report/export/png")
def export_report_png(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")):
    """Export report as PNG (single page image)"""
    from PIL import Image, ImageDraw, ImageFont
    import base64
    import tempfile
    
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    items = REPORT_STORE[x_report_id]
    
    # Create a white background image
    width, height = 2480, 3508  # A4 size at 300 DPI
    img = Image.new('RGB', (width, height), 'white')
    draw = ImageDraw.Draw(img)
    
    try:
        # Try to use a nice font
        try:
            font_title = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 60)
            font_text = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 40)
        except:
            font_title = ImageFont.load_default()
            font_text = ImageFont.load_default()
        
        y_offset = 100
        
        # Add title
        draw.text((100, y_offset), "Experiment Report", fill='black', font=font_title)
        y_offset += 150
        
        # Add each item
        for idx, item in enumerate(items):
            # Item title
            draw.text((100, y_offset), f"{idx + 1}. {item['title']}", fill='black', font=font_text)
            y_offset += 80
            
            # Add image if present
            if item['type'] in ['chart', 'table'] and 'imageDataUrl' in item.get('content', {}):
                try:
                    img_data = item['content']['imageDataUrl']
                    if img_data.startswith('data:image'):
                        img_data = img_data.split(',')[1]
                    
                    item_img_bytes = base64.b64decode(img_data)
                    item_img = Image.open(io.BytesIO(item_img_bytes))
                    
                    # Resize to fit
                    max_width = width - 200
                    max_height = 600
                    # Use Resampling.LANCZOS if available, otherwise fall back to LANCZOS
                    try:
                        item_img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                    except AttributeError:
                        item_img.thumbnail((max_width, max_height), Image.LANCZOS)
                    
                    # Paste image
                    img.paste(item_img, (100, y_offset))
                    y_offset += item_img.height + 50
                except Exception as e:
                    print(f"Error adding image: {e}")
            
            # Add text
            if item['type'] == 'text' and 'text' in item.get('content', {}):
                text = item['content']['text'][:200]  # Limit text length
                draw.text((100, y_offset), text, fill='black', font=font_text)
                y_offset += 100
            
            # Add comment if present
            if item.get('comment'):
                comment_y = y_offset
                # Calculate comment box height based on text length
                comment_text = f"Comment: {item['comment'][:150]}"
                # Estimate height: ~40 pixels per line, max 3 lines
                lines = comment_text.split('\n')
                num_lines = min(len(lines), 3)
                comment_box_height = max(60, num_lines * 40)
                
                # Draw comment box background (yellow highlight)
                draw.rectangle([90, comment_y - 5, width - 100, comment_y + comment_box_height], fill='#FFFACD', outline='#FFD700', width=2)
                
                # Draw comment text (wrap if needed)
                try:
                    # Try to use textbbox if available (Pillow >= 8.0)
                    if hasattr(draw, 'textbbox'):
                        bbox = draw.textbbox((100, comment_y), comment_text, font=font_text)
                        actual_height = bbox[3] - bbox[1] + 20
                        comment_box_height = max(comment_box_height, actual_height)
                except:
                    pass
                
                # Draw comment text
                draw.text((100, comment_y), comment_text, fill='#333333', font=font_text)
                y_offset += comment_box_height + 20
            
            y_offset += 100
            
            if y_offset > height - 200:
                break  # Stop if we run out of space
        
    except Exception as e:
        print(f"Error creating PNG: {e}")
    
    # Save to bytes
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    
    return Response(
        content=buffer.read(),
        media_type="image/png",
        headers={"Content-Disposition": f"attachment; filename=experiment_report_{pd.Timestamp.now().strftime('%Y%m%d')}.png"}
    )


@app.get("/report/export/word")
def export_report_word(x_report_id: Optional[str] = Header(default=None, alias="x-report-id")):
    """Export report as Word document"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import base64
        import tempfile
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"Word export library not available: {str(e)}")
    
    if not x_report_id or x_report_id not in REPORT_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing report_id")
    
    try:
        items = REPORT_STORE[x_report_id]
        
        # Create document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add title
        title = doc.add_heading('Experiment Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add each item
        for idx, item in enumerate(items):
            # Item title
            doc.add_heading(f"{idx + 1}. {item['title']}", level=1)
            
            # Add image if present
            if item['type'] in ['chart', 'table'] and 'imageDataUrl' in item.get('content', {}):
                try:
                    img_data = item['content']['imageDataUrl']
                    if img_data.startswith('data:image'):
                        img_data = img_data.split(',')[1]
                    
                    img_bytes = base64.b64decode(img_data)
                    img_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    img_file.write(img_bytes)
                    img_file.close()
                    
                    # Add image to document
                    doc.add_picture(img_file.name, width=Inches(6))
                    
                    # Clean up
                    os.unlink(img_file.name)
                except Exception as e:
                    print(f"Error adding image to Word: {e}")
            
            # Add text content
            if item['type'] == 'text' and 'text' in item.get('content', {}):
                doc.add_paragraph(item['content']['text'])
            
            # Add comment if present
            if item.get('comment'):
                para = doc.add_paragraph(f"Comment: {item['comment']}")
                para.italic = True
            
            # Add spacing
            doc.add_paragraph()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return Response(
            content=buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=experiment_report_{pd.Timestamp.now().strftime('%Y%m%d')}.docx",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "x-report-id",
            }
        )
    except Exception as e:
        print(f"Error generating Word document: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")


# =============================================================================
# METRIC FUNCTIONS API
# =============================================================================

@app.post("/functions/test", response_model=FunctionTestResponse)
async def test_metric_function(request: FunctionTestRequest) -> FunctionTestResponse:
    """
    Test a metric function in sandbox mode.
    Returns a preview of the output (limited rows).
    """
    try:
        result = test_function(
            code=request.code,
            parameters=request.parameters,
            username=request.username,
            limit_rows=100
        )
        return FunctionTestResponse(**result)
    except Exception as e:
        return FunctionTestResponse(
            success=False,
            error=f"Test failed: {str(e)}",
            row_count=0
        )


@app.post("/functions/execute", response_model=FunctionExecuteResponse)
async def execute_metric_function(request: FunctionExecuteRequest) -> FunctionExecuteResponse:
    """
    Execute a metric function and return full results.
    """
    try:
        result_df, error, output_columns = execute_function(
            code=request.code,
            parameters=request.parameters,
            username=request.username
        )
        
        if error:
            return FunctionExecuteResponse(
                success=False,
                error=error,
                row_count=0
            )
        
        return FunctionExecuteResponse(
            success=True,
            data=result_df.to_dict(orient='records'),
            columns=list(result_df.columns),
            output_columns=output_columns,
            row_count=len(result_df)
        )
    except Exception as e:
        return FunctionExecuteResponse(
            success=False,
            error=f"Execution failed: {str(e)}",
            row_count=0
        )


@app.post("/functions/preview", response_model=FunctionPreviewResponse)
async def preview_function_result(
    request: FunctionPreviewRequest
) -> FunctionPreviewResponse:
    """
    Execute a function and return a preview of results with statistics.
    Does NOT join with CSV - just shows the function output.
    """
    try:
        # Execute the function
        result_df, error, output_columns = execute_function(
            code=request.code,
            parameters=request.parameters,
            username=request.username
        )
        
        if error:
            return FunctionPreviewResponse(
                success=False,
                error=error,
                row_count=0
            )
        
        # Calculate descriptive stats for all columns
        stats = {}
        for col in result_df.columns:
            col_data = result_df[col]
            if pd.api.types.is_numeric_dtype(col_data):
                stats[col] = {
                    'type': 'numeric',
                    'count': int(col_data.notna().sum()),
                    'mean': round(float(col_data.mean()), 2) if col_data.notna().any() else None,
                    'std': round(float(col_data.std()), 2) if col_data.notna().any() else None,
                    'min': round(float(col_data.min()), 2) if col_data.notna().any() else None,
                    'max': round(float(col_data.max()), 2) if col_data.notna().any() else None,
                    'median': round(float(col_data.median()), 2) if col_data.notna().any() else None,
                    'null_count': int(col_data.isna().sum()),
                }
            else:
                stats[col] = {
                    'type': 'categorical',
                    'count': int(col_data.notna().sum()),
                    'unique': int(col_data.nunique()),
                    'null_count': int(col_data.isna().sum()),
                    'top_values': {str(k): int(v) for k, v in col_data.value_counts().head(5).items()} if col_data.notna().any() else {},
                }
        
        # Prepare preview data (first 100 rows)
        preview_df = result_df.head(100).copy()
        for col in preview_df.columns:
            if pd.api.types.is_datetime64_any_dtype(preview_df[col]):
                preview_df[col] = preview_df[col].dt.strftime('%Y-%m-%d')
        preview_data = preview_df.fillna('').to_dict(orient='records')
        
        return FunctionPreviewResponse(
            success=True,
            preview=preview_data,
            columns=list(result_df.columns),
            row_count=len(result_df),
            stats=stats
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return FunctionPreviewResponse(
            success=False,
            error=f"Preview failed: {str(e)}",
            row_count=0
        )


@app.post("/functions/join", response_model=FunctionJoinResponse)
async def join_function_with_csv(
    request: FunctionJoinRequest,
    x_session_id: Optional[str] = Header(default=None)
) -> FunctionJoinResponse:
    """
    Execute a function and join its results with the uploaded CSV.
    Updates the session DataFrame with the new columns.
    Supports configurable join columns and join type.
    """
    if not x_session_id or x_session_id not in SESSION_STORE:
        return FunctionJoinResponse(
            success=False,
            error="Invalid or missing session_id. Upload data first.",
            row_count=0,
            matched_rows=0
        )
    
    try:
        session = SESSION_STORE[x_session_id]
        was_duckdb = isinstance(session, dict) and "parquet_path" in session
        parquet_path_for_update = session.get("parquet_path") if was_duckdb else None
        
        # DuckDB sessions store dict with parquet_path; load DataFrame for join
        if was_duckdb:
            parquet_path = session["parquet_path"]
            if not os.path.exists(parquet_path):
                return FunctionJoinResponse(
                    success=False,
                    error="Session data file not found.",
                    row_count=0,
                    matched_rows=0
                )
            con = duckdb.connect()
            csv_df = con.execute(f"SELECT * FROM read_parquet('{parquet_path}')").fetchdf()
            con.close()
        else:
            csv_df = session.copy()
        
        # Execute the function
        result_df, error, output_columns = execute_function(
            code=request.code,
            parameters=request.parameters,
            username=request.username
        )
        
        if error:
            return FunctionJoinResponse(
                success=False,
                error=error,
                row_count=0,
                matched_rows=0
            )
        
        # Remove any duplicate columns in CSV
        csv_df = csv_df.loc[:, ~csv_df.columns.duplicated()]
        result_df = result_df.loc[:, ~result_df.columns.duplicated()]
        
        # Prepare CSV for join - create yyyymmdd column if not exists
        if 'yyyymmdd' not in csv_df.columns:
            if 'time' in csv_df.columns:
                # time is already in YYYYMMDD format (as int or string)
                csv_df['yyyymmdd'] = csv_df['time'].astype(str)
            elif 'date' in csv_df.columns:
                # date is in datetime format, convert to YYYYMMDD string
                if pd.api.types.is_datetime64_any_dtype(csv_df['date']):
                    csv_df['yyyymmdd'] = csv_df['date'].dt.strftime('%Y%m%d')
                else:
                    # Try parsing as date
                    csv_df['yyyymmdd'] = pd.to_datetime(csv_df['date']).dt.strftime('%Y%m%d')
        
        # Handle captain_id column name variations
        if 'captain_id' not in csv_df.columns:
            if 'captainId' in csv_df.columns:
                csv_df['captain_id'] = csv_df['captainId'].astype(str)
            elif 'captainid' in csv_df.columns:
                csv_df['captain_id'] = csv_df['captainid'].astype(str)
        
        # Ensure captain_id is string in both dataframes
        if 'captain_id' in csv_df.columns:
            csv_df['captain_id'] = csv_df['captain_id'].astype(str)
        if 'captain_id' in result_df.columns:
            result_df['captain_id'] = result_df['captain_id'].astype(str)
        
        # Ensure yyyymmdd is string in both dataframes  
        if 'yyyymmdd' in csv_df.columns:
            csv_df['yyyymmdd'] = csv_df['yyyymmdd'].astype(str)
        if 'yyyymmdd' in result_df.columns:
            result_df['yyyymmdd'] = result_df['yyyymmdd'].astype(str)
        
        # Use join columns from request (default: ['captain_id', 'yyyymmdd'])
        join_columns = request.join_columns
        join_type = request.join_type if request.join_type in ['left', 'inner'] else 'left'
        
        # Check for duplicate join key combinations in function result
        if result_df.duplicated(subset=join_columns).any():
            dup_count = result_df.duplicated(subset=join_columns).sum()
            return FunctionJoinResponse(
                success=False,
                error=f"Function result has {dup_count} duplicate {' + '.join(join_columns)} combinations. Each combination must be unique. Please fix the function to return unique rows.",
                row_count=0,
                matched_rows=0
            )
        
        # Debug: log unique values for joining
        print(f"DEBUG: Join columns: {join_columns}, Join type: {join_type}")
        print(f"DEBUG: CSV columns: {list(csv_df.columns)}")
        print(f"DEBUG: Result columns: {list(result_df.columns)}")
        
        # Verify join columns exist
        for col in join_columns:
            if col not in csv_df.columns:
                return FunctionJoinResponse(
                    success=False,
                    error=f"CSV is missing required join column: '{col}'. Available columns: {list(csv_df.columns)}",
                    row_count=0,
                    matched_rows=0
                )
            if col not in result_df.columns:
                return FunctionJoinResponse(
                    success=False,
                    error=f"Function result is missing required join column: '{col}'. Available columns: {list(result_df.columns)}",
                    row_count=0,
                    matched_rows=0
                )
        
        # Perform the join
        merged_df = csv_df.merge(
            result_df,
            on=join_columns,
            how=join_type,
            suffixes=('', '_computed')
        )
        
        print(f"DEBUG: Merged from {len(csv_df)} CSV rows + {len(result_df)} result rows = {len(merged_df)} merged rows")
        
        # Count matched rows (rows with non-null values in new columns)
        matched_rows = 0
        if output_columns and output_columns[0] in merged_df.columns:
            matched_rows = merged_df[output_columns[0]].notna().sum()
        
        print(f"DEBUG: Matched rows: {matched_rows} out of {len(merged_df)}")
        
        # Calculate descriptive stats for new columns
        metrics_stats = {}
        for col in output_columns or []:
            if col in merged_df.columns:
                col_data = merged_df[col]
                if pd.api.types.is_numeric_dtype(col_data):
                    stats = {
                        'count': int(col_data.notna().sum()),
                        'mean': round(float(col_data.mean()), 2) if col_data.notna().any() else None,
                        'std': round(float(col_data.std()), 2) if col_data.notna().any() else None,
                        'min': round(float(col_data.min()), 2) if col_data.notna().any() else None,
                        'max': round(float(col_data.max()), 2) if col_data.notna().any() else None,
                        'median': round(float(col_data.median()), 2) if col_data.notna().any() else None,
                        'null_count': int(col_data.isna().sum()),
                    }
                else:
                    # For non-numeric columns
                    stats = {
                        'count': int(col_data.notna().sum()),
                        'unique': int(col_data.nunique()),
                        'null_count': int(col_data.isna().sum()),
                        'top_values': col_data.value_counts().head(5).to_dict() if col_data.notna().any() else {},
                    }
                metrics_stats[col] = stats
        
        # Update the session store: for DuckDB sessions, write merged data back to parquet and keep DuckDB session
        if was_duckdb and parquet_path_for_update:
            import pyarrow as pa
            import pyarrow.parquet as pq
            table = pa.Table.from_pandas(merged_df, preserve_index=False)
            pq.write_table(table, parquet_path_for_update)
            # Build metadata from merged_df for DuckDB session
            cols = list(merged_df.columns)
            cohorts = sorted(merged_df["cohort"].dropna().unique().tolist()) if "cohort" in merged_df.columns else []
            date_min = merged_df["date"].min()
            date_max = merged_df["date"].max()
            if hasattr(date_min, "strftime"):
                date_min_str = date_min.strftime("%Y-%m-%d")
                date_max_str = date_max.strftime("%Y-%m-%d")
            else:
                date_min_str = str(date_min)[:10] if date_min else None
                date_max_str = str(date_max)[:10] if date_max else None
            metrics = [c for c in cols if c not in {"cohort", "date", "time"}]
            SESSION_STORE[x_session_id] = {
                "parquet_path": parquet_path_for_update,
                "num_rows": len(merged_df),
                "columns": cols,
                "cohorts": cohorts,
                "date_min": date_min_str,
                "date_max": date_max_str,
                "metrics": metrics,
            }
        else:
            SESSION_STORE[x_session_id] = merged_df
        
        # Prepare preview data (first 100 rows)
        preview_df = merged_df.head(100).copy()
        # Convert date columns to string for JSON serialization
        for col in preview_df.columns:
            if pd.api.types.is_datetime64_any_dtype(preview_df[col]):
                preview_df[col] = preview_df[col].dt.strftime('%Y-%m-%d')
        preview_data = preview_df.fillna('').to_dict(orient='records')
        
        return FunctionJoinResponse(
            success=True,
            added_columns=output_columns,
            row_count=len(merged_df),
            matched_rows=int(matched_rows),
            preview=preview_data,
            columns=list(merged_df.columns),
            metrics_stats=metrics_stats
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return FunctionJoinResponse(
            success=False,
            error=f"Join failed: {str(e)}",
            row_count=0,
            matched_rows=0
        )


def _stream_parquet_as_csv(parquet_path: str, chunk_size: int = 50000):
    """Yield CSV chunks from parquet via DuckDB without loading full data into memory."""
    con = duckdb.connect()
    try:
        first = True
        offset = 0
        while True:
            df_chunk = con.execute(
                f"SELECT * FROM read_parquet('{parquet_path}') LIMIT {chunk_size} OFFSET {offset}"
            ).fetchdf()
            if df_chunk is None or len(df_chunk) == 0:
                break
            buf = io.StringIO()
            df_chunk.to_csv(buf, index=False, header=first)
            first = False
            yield buf.getvalue()
            if len(df_chunk) < chunk_size:
                break
            offset += chunk_size
    finally:
        con.close()


@app.get("/data/download")
async def download_session_data(
    x_session_id: Optional[str] = Header(default=None)
):
    """
    Download the current session data as a CSV file.
    For DuckDB sessions: streams CSV from parquet (fast, no full load).
    For in-memory sessions: returns CSV from DataFrame.
    """
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    
    session = SESSION_STORE[x_session_id]
    
    # DuckDB session: stream CSV from parquet in chunks
    if isinstance(session, dict) and "parquet_path" in session:
        parquet_path = session["parquet_path"]
        if not os.path.exists(parquet_path):
            raise HTTPException(status_code=400, detail="Session data file not found.")
        return StreamingResponse(
            _stream_parquet_as_csv(parquet_path),
            media_type="text/csv",
            headers={
                "Content-Disposition": "attachment; filename=data_with_metrics.csv",
                "Cache-Control": "no-cache",
            },
        )
    
    df = session
    csv_buffer = io.StringIO()
    df.to_csv(csv_buffer, index=False)
    csv_content = csv_buffer.getvalue()
    return Response(
        content=csv_content,
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=data_with_metrics.csv"},
    )


@app.get("/data/preview")
async def preview_session_data(
    x_session_id: Optional[str] = Header(default=None),
    limit: int = 100
):
    """
    Get a preview of the current session data.
    """
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    
    session = SESSION_STORE[x_session_id]
    if isinstance(session, dict) and "parquet_path" in session:
        parquet_path = session["parquet_path"]
        if not os.path.exists(parquet_path):
            raise HTTPException(status_code=400, detail="Session data file not found.")
        con = duckdb.connect()
        preview_df = con.execute(
            f"SELECT * FROM read_parquet('{parquet_path}') LIMIT {limit}"
        ).fetchdf()
        con.close()
        total_rows = session["num_rows"]
        columns = session["columns"]
    else:
        df = session
        preview_df = df.head(limit).copy()
        total_rows = len(df)
        columns = list(df.columns)
    
    for col in preview_df.columns:
        if pd.api.types.is_datetime64_any_dtype(preview_df[col]):
            preview_df[col] = preview_df[col].dt.strftime('%Y-%m-%d')
    
    return {
        "columns": columns,
        "preview": preview_df.fillna('').to_dict(orient='records'),
        "total_rows": total_rows,
    }


def _quote_col(col: str) -> str:
    """Quote column name for DuckDB SQL."""
    return f'"{col.replace(chr(34), chr(34) + chr(34))}"'


def compute_pivot_duckdb(parquet_path: str, payload: PivotRequest) -> PivotResponse:
    """
    Run pivot aggregation on Parquet via DuckDB.
    Returns columns + data matching frontend PivotResult (wide table: one row per row_key, value columns per col_key × value spec).
    """
    if not os.path.exists(parquet_path):
        raise HTTPException(status_code=400, detail="Session data file not found")

    con = duckdb.connect()
    try:
        desc = con.execute(f"DESCRIBE SELECT * FROM read_parquet('{parquet_path}')").fetchdf()
        name_col = desc.columns[0]
        type_col = desc.columns[1]
        col_types = dict(zip(desc[name_col], desc[type_col]))
        all_cols = set(col_types.keys())

        def _is_numeric_col(col: str) -> bool:
            t = col_types.get(col, "")
            t_lower = str(t).lower()
            return any(x in t_lower for x in ["int", "float", "double", "decimal", "numeric", "bigint", "smallint"])

        row_fields = [c for c in payload.row_fields if c in all_cols]
        col_fields = [c for c in payload.col_fields if c in all_cols]
        values = [v for v in payload.values if v.col in all_cols]
        if not values:
            return PivotResponse(columns=row_fields, data=[])

        # Build WHERE from filters
        where_parts = []
        for f in payload.filters:
            if f.column not in all_cols:
                continue
            q = _quote_col(f.column)
            op = f.operator
            val = f.value
            if op == "equals":
                if isinstance(val, str):
                    where_parts.append(f"{q} = '{str(val).replace(chr(39), chr(39)+chr(39))}'")
                else:
                    where_parts.append(f"{q} = {val}")
            elif op == "not_equals":
                if isinstance(val, str):
                    where_parts.append(f"{q} <> '{str(val).replace(chr(39), chr(39)+chr(39))}'")
                else:
                    where_parts.append(f"{q} <> {val}")
            elif op == "contains":
                where_parts.append(f"LOWER(CAST({q} AS VARCHAR)) LIKE '%{str(val).lower().replace(chr(39), chr(39)+chr(39))}%'")
            elif op == "not_contains":
                where_parts.append(f"LOWER(CAST({q} AS VARCHAR)) NOT LIKE '%{str(val).lower().replace(chr(39), chr(39)+chr(39))}%'")
            elif op == "in":
                if isinstance(val, list):
                    in_vals = []
                    for x in val:
                        if isinstance(x, str):
                            in_vals.append(f"'{str(x).replace(chr(39), chr(39)+chr(39))}'")
                        else:
                            in_vals.append(str(x))
                    where_parts.append(f"{q} IN ({','.join(in_vals)})")
            elif op == "between" and isinstance(val, (list, tuple)) and len(val) >= 2:
                a, b = val[0], val[1]
                if isinstance(a, str):
                    where_parts.append(f"{q} >= '{str(a).replace(chr(39), chr(39)+chr(39))}' AND {q} <= '{str(b).replace(chr(39), chr(39)+chr(39))}'")
                else:
                    where_parts.append(f"{q} >= {a} AND {q} <= {b}")

        where_sql = " AND ".join(where_parts) if where_parts else "1=1"

        # SELECT: row fields, col fields, value aggregates (non-numeric cols use COUNT/MIN/MAX on raw value)
        select_parts = [_quote_col(c) for c in row_fields + col_fields]
        value_aliases = []
        for v in values:
            q = _quote_col(v.col)
            agg = v.agg
            numeric = _is_numeric_col(v.col)
            if agg == "sum":
                if numeric:
                    select_parts.append(f"SUM(CAST({q} AS DOUBLE)) AS {v.col}__{agg}")
                else:
                    select_parts.append(f"COUNT({q}) AS {v.col}__{agg}")
            elif agg == "avg":
                if numeric:
                    select_parts.append(f"AVG(CAST({q} AS DOUBLE)) AS {v.col}__{agg}")
                else:
                    select_parts.append(f"COUNT({q}) AS {v.col}__{agg}")
            elif agg == "min":
                if numeric:
                    select_parts.append(f"MIN(CAST({q} AS DOUBLE)) AS {v.col}__{agg}")
                else:
                    select_parts.append(f"COUNT({q}) AS {v.col}__{agg}")
            elif agg == "max":
                if numeric:
                    select_parts.append(f"MAX(CAST({q} AS DOUBLE)) AS {v.col}__{agg}")
                else:
                    select_parts.append(f"COUNT({q}) AS {v.col}__{agg}")
            elif agg == "count":
                select_parts.append(f"COUNT({q}) AS {v.col}__{agg}")
            elif agg == "countDistinct":
                select_parts.append(f"COUNT(DISTINCT {q}) AS {v.col}__{agg}")
            else:
                agg = "sum"
                if numeric:
                    select_parts.append(f"SUM(CAST({q} AS DOUBLE)) AS {v.col}__{agg}")
                else:
                    select_parts.append(f"COUNT({q}) AS {v.col}__{agg}")
            value_aliases.append(f"{v.col}__{agg}")

        group_cols = row_fields + col_fields
        if group_cols:
            group_by = ", ".join([_quote_col(c) for c in group_cols])
            sql = f"""
            SELECT {", ".join(select_parts)}
            FROM read_parquet('{parquet_path}')
            WHERE {where_sql}
            GROUP BY {group_by}
            """
        else:
            sql = f"""
            SELECT {", ".join(select_parts)}
            FROM read_parquet('{parquet_path}')
            WHERE {where_sql}
            """
        df = con.execute(sql).fetchdf()
    finally:
        con.close()

    if df.empty:
        out_columns = row_fields + [f"{v.col} ({v.agg})" for v in values] if not col_fields else row_fields
        return PivotResponse(columns=out_columns, data=[])

    def _clean_val(v):
        if v is None:
            return None
        if isinstance(v, float) and (pd.isna(v) or v == float("inf") or v == float("-inf")):
            return None
        if isinstance(v, (pd.Timestamp,)):
            return v.strftime("%Y-%m-%d")
        return v

    # Reshape to wide table matching frontend PivotResult
    def display_col_key(parts):
        s = "¦".join("" if p is None or (isinstance(p, float) and pd.isna(p)) else str(p) for p in parts)
        return (s.strip() or "All").replace("¦", " / ")

    row_cols = row_fields
    col_cols = col_fields
    sorted_row_keys = []
    row_key_to_idx = {}
    for _, r in df.iterrows():
        key = tuple(r[c] for c in row_cols)
        if key not in row_key_to_idx:
            row_key_to_idx[key] = len(sorted_row_keys)
            sorted_row_keys.append(key)

    col_keys_set = set()
    for _, r in df.iterrows():
        key = tuple(r[c] for c in col_cols) if col_cols else ("",)
        col_keys_set.add(key)
    sorted_col_keys = sorted(col_keys_set, key=lambda k: "¦".join("" if x is None or (isinstance(x, float) and pd.isna(x)) else str(x) for x in k))

    if not col_cols:
        out_columns = list(row_cols) + [f"{v.col} ({v.agg})" for v in values]
        data = []
        for rk in sorted_row_keys:
            row_dict = {}
            for i, c in enumerate(row_cols):
                val = rk[i] if i < len(rk) else None
                if isinstance(val, (pd.Timestamp,)):
                    val = val.strftime("%Y-%m-%d") if pd.notna(val) else None
                row_dict[c] = val
            sub = df
            for i, c in enumerate(row_cols):
                vv = rk[i] if i < len(rk) else None
                if vv is None or (isinstance(vv, float) and pd.isna(vv)):
                    sub = sub[sub[c].isna()]
                else:
                    sub = sub[sub[c] == vv]
            for v in values:
                alias = f"{v.col}__{v.agg}"
                col_name = f"{v.col} ({v.agg})"
                row_dict[col_name] = float(sub[alias].iloc[0]) if len(sub) and alias in sub.columns else None
            data.append(row_dict)
        for row in data:
            for k in list(row.keys()):
                row[k] = _clean_val(row[k])
        return PivotResponse(columns=out_columns, data=data)

    out_columns = list(row_cols) + [
        f"{display_col_key(ck)} • {v.col} ({v.agg})"
        for ck in sorted_col_keys
        for v in values
    ]
    data = []
    for rk in sorted_row_keys:
        row_dict = {}
        for i, c in enumerate(row_cols):
            val = rk[i] if i < len(rk) else None
            if isinstance(val, (pd.Timestamp,)):
                val = val.strftime("%Y-%m-%d") if pd.notna(val) else None
            row_dict[c] = val
        sub = df
        for i, c in enumerate(row_cols):
            vv = rk[i] if i < len(rk) else None
            if vv is None or (isinstance(vv, float) and pd.isna(vv)):
                sub = sub[sub[c].isna()]
            else:
                sub = sub[sub[c] == vv]
        for ck in sorted_col_keys:
            sub2 = sub
            for i, c in enumerate(col_cols):
                vv = ck[i] if i < len(ck) else None
                if vv is None or (isinstance(vv, float) and pd.isna(vv)):
                    sub2 = sub2[sub2[c].isna()]
                else:
                    sub2 = sub2[sub2[c] == vv]
            for v in values:
                alias = f"{v.col}__{v.agg}"
                col_name = f"{display_col_key(ck)} • {v.col} ({v.agg})"
                row_dict[col_name] = float(sub2[alias].iloc[0]) if len(sub2) and alias in sub2.columns else None
        data.append(row_dict)

    for row in data:
        for k in list(row.keys()):
            row[k] = _clean_val(row[k])

    return PivotResponse(columns=out_columns, data=data)


@app.post("/pivot", response_model=PivotResponse, responses={400: {"model": ErrorResponse}})
def pivot(
    payload: PivotRequest,
    x_session_id: Optional[str] = Header(default=None),
) -> PivotResponse:
    """
    Run pivot aggregation on session data.
    For DuckDB (large file) sessions: runs GROUP BY + aggregates in DuckDB and returns wide table.
    Requires x-session-id header.
    """
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    session = SESSION_STORE[x_session_id]
    if not isinstance(session, dict) or "parquet_path" not in session:
        raise HTTPException(status_code=400, detail="Pivot is only supported for large (DuckDB) sessions. Use the Pivot tab with an uploaded large file.")
    parquet_path = session["parquet_path"]
    return compute_pivot_duckdb(parquet_path, payload)


@app.get("/data/session")
async def get_session_data(
    x_session_id: Optional[str] = Header(default=None)
):
    """
    Get session data as JSON for frontend state refresh.
    For DuckDB (large file) sessions: returns metadata + empty rows (analysis via backend).
    For in-memory sessions: returns metadata + full rows, capped at 2000 for response size.
    """
    import numpy as np
    
    if not x_session_id or x_session_id not in SESSION_STORE:
        raise HTTPException(status_code=400, detail="Invalid or missing session_id. Upload data first.")
    
    session = SESSION_STORE[x_session_id]
    
    # DuckDB session: return metadata only, no full rows (fast response; analysis uses /insights)
    if isinstance(session, dict) and "parquet_path" in session:
        columns = session["columns"]
        metric_columns = [c for c in columns if c not in {"cohort", "date", "time"}]
        return {
            "rows": [],
            "columns": columns,
            "metric_columns": metric_columns,
            "numeric_columns": session.get("numeric_columns", session.get("metrics", [])),
            "categorical_columns": session.get("categorical_columns", []),
            "cohorts": session["cohorts"],
            "date_min": session.get("date_min"),
            "date_max": session.get("date_max"),
            "row_count": session["num_rows"],
        }
    
    df = session
    # Cap rows for in-memory sessions to keep response fast (frontend can use /insights for full analysis)
    max_rows = 2000
    df_slice = df.head(max_rows)
    
    def clean_value(v):
        if v is None:
            return None
        if isinstance(v, float):
            if np.isnan(v) or np.isinf(v):
                return None
        return v
    
    df_copy = df_slice.copy()
    for col in df_copy.columns:
        if pd.api.types.is_datetime64_any_dtype(df_copy[col]):
            df_copy[col] = df_copy[col].dt.strftime('%Y-%m-%d')
    records = df_copy.to_dict(orient='records')
    cleaned_records = [{k: clean_value(v) for k, v in r.items()} for r in records]
    
    columns = list(df.columns.astype(str))
    metric_columns = [c for c in columns if c not in {"cohort", "date", "time"}]
    numeric_columns = [c for c in metric_columns if pd.api.types.is_numeric_dtype(df[c])]
    categorical_columns = []
    for col in columns:
        if col not in {"cohort", "date", "time"}:
            if df[col].dtype == 'object' or df[col].dtype.name == 'category':
                categorical_columns.append(col)
            elif col in numeric_columns:
                n = df[col].nunique()
                if n < 20 and n < len(df) * 0.1:
                    categorical_columns.append(col)
    
    cohorts = sorted(df["cohort"].dropna().unique().tolist()) if "cohort" in df.columns else []
    date_min = df["date"].min().strftime("%Y-%m-%d") if "date" in df.columns and pd.notna(df["date"].min()) else None
    date_max = df["date"].max().strftime("%Y-%m-%d") if "date" in df.columns and pd.notna(df["date"].max()) else None
    
    return {
        "rows": cleaned_records,
        "columns": columns,
        "metric_columns": metric_columns,
        "numeric_columns": numeric_columns,
        "categorical_columns": sorted(categorical_columns),
        "cohorts": cohorts,
        "date_min": date_min,
        "date_max": date_max,
        "row_count": len(df),
    }


@app.get("/functions/template", response_model=FunctionTemplateResponse)
async def get_function_template() -> FunctionTemplateResponse:
    """
    Get the template code for creating a new metric function.
    """
    return FunctionTemplateResponse(template=FUNCTION_TEMPLATE)


@app.get("/health")
def health():
    """Health check endpoint with system status."""
    import psutil
    
    # Get memory info
    memory = psutil.virtual_memory()
    disk = psutil.disk_usage('/')
    
    return {
        "status": "ok",
        "active_sessions": len(SESSION_STORE),
        "active_uploads": len(UPLOAD_PROGRESS),
        "memory_percent": memory.percent,
        "disk_percent": disk.percent,
        "max_file_size_gb": MAX_FILE_SIZE / (1024**3),
    }


@app.on_event("startup")
async def startup_event():
    """Clean up any orphaned temp files on startup."""
    import logging
    logging.info("Starting application with 5GB file upload support")
    
    # Clean up any orphaned temp directories from previous runs
    temp_base = tempfile.gettempdir()
    cleaned_count = 0
    try:
        for item in os.listdir(temp_base):
            item_path = os.path.join(temp_base, item)
            # Only clean up directories that look like our temp dirs
            if os.path.isdir(item_path) and item.startswith('tmp'):
                try:
                    # Check if it contains our upload file
                    upload_file = os.path.join(item_path, 'upload.csv')
                    if os.path.exists(upload_file):
                        shutil.rmtree(item_path, ignore_errors=True)
                        cleaned_count += 1
                except Exception:
                    pass
    except Exception as e:
        logging.warning(f"Temp cleanup failed: {e}")
    
    if cleaned_count > 0:
        logging.info(f"Cleaned up {cleaned_count} orphaned temp directories")


@app.on_event("shutdown")
async def shutdown_event():
    """Clean up temp files on shutdown."""
    import logging
    logging.info("Shutting down, cleaning up temp files...")
    
    # Clean up DuckDB session parquet files
    for session_id, session in list(SESSION_STORE.items()):
        try:
            if isinstance(session, dict) and "parquet_path" in session:
                parquet_path = session.get("parquet_path")
                if parquet_path and os.path.exists(parquet_path):
                    os.unlink(parquet_path)
                temp_dir = os.path.dirname(parquet_path) if parquet_path else None
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass
    
    # Clean up all session temp files
    for session_id, temp_path in list(SESSION_FILE_STORE.items()):
        try:
            if os.path.isdir(temp_path):
                shutil.rmtree(temp_path, ignore_errors=True)
            else:
                temp_dir = os.path.dirname(temp_path)
                if temp_dir and os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass
    
    # Clean up pending uploads
    for upload_id, progress in list(UPLOAD_PROGRESS.items()):
        try:
            temp_dir = progress.get("temp_dir")
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass


if __name__ == "__main__":
    import uvicorn
    # Render sets PORT environment variable, default to 8000 for local development
    port = int(os.environ.get("PORT", "8001"))
    
    # Configure uvicorn with extended timeouts for large file uploads (5GB support)
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=port,
        timeout_keep_alive=300,  # 5 minutes keep-alive timeout
        # Note: For production, consider setting these via environment variables:
        # - limit_concurrency: limits concurrent connections
        # - limit_max_requests: limits requests per worker
    )

